from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(
    r"D:\School materials\SWP391\AI-student-hub\output\test-plan"
    r"\AI_StudyHub_Comprehensive_Test_Plan_SRS_v3.3.docx"
)
SRS_PATH = (
    r"D:\School materials\SWP391\FinalReport\SRS"
    r"\AI_StudyHub_SRS_v3.3_Visual_Paradigm_Use_Cases.docx"
)

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "6B7280"
WHITE = "FFFFFF"
GREEN = "DFF2E1"
AMBER = "FFF3CD"
RED = "FDE2E2"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[index]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_run_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = "Calibri"
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_font(paragraph, size=10.5, color=INK):
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color)


def add_label_value(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    label_run = p.add_run(f"{label}: ")
    set_run_font(label_run, size=10.5, bold=True, color=DARK_BLUE)
    value_run = p.add_run(value)
    set_run_font(value_run, size=10.5, color=INK)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=INK)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=INK)
    return p


def add_callout(doc, label, text, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    lead = p.add_run(f"{label}: ")
    set_run_font(lead, size=10.5, bold=True, color=DARK_BLUE)
    body = p.add_run(text)
    set_run_font(body, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


def case(
    uc_id,
    title,
    actor,
    preconditions,
    data,
    steps,
    expected,
    negative,
    rules,
    priority="High",
):
    return {
        "id": uc_id,
        "title": title,
        "actor": actor,
        "preconditions": preconditions,
        "data": data,
        "steps": steps,
        "expected": expected,
        "negative": negative,
        "rules": rules,
        "priority": priority,
    }


GROUPS = [
    (
        "A. Public Access",
        [
            case(
                "UC-PUB-01",
                "Browse Public Libraries",
                "Guest / Anonymous User",
                "At least one public and one private library exist; public library contains APPROVED and non-approved documents.",
                "PUB_LIB_01 public; PRIV_LIB_01 private; documents in APPROVED, PENDING and deleted states.",
                [
                    "Open the landing page without an authenticated session.",
                    "Choose Access as Guest and open Discover or Libraries.",
                    "Call GET /api/public/libraries and inspect returned metadata.",
                    "Open PUB_LIB_01 and inspect the listed documents.",
                    "Search or filter the public catalog if controls are available.",
                ],
                "Only public libraries are returned. Document counts and detail contain only public, APPROVED, non-deleted documents. No Authorization header is required.",
                [
                    "Verify PRIV_LIB_01 cannot be opened through the public endpoint.",
                    "Verify PENDING, FLAGGED, REJECTED and deleted documents do not appear.",
                    "Verify Guest does not receive edit, upload or delete controls.",
                ],
                "FR-PUB-01; FR-PUB-02; BR-DOC-07; BR-SEC-03",
            ),
            case(
                "UC-PUB-02",
                "Download Public Document",
                "Guest / Anonymous User",
                "A public library contains a public APPROVED document with a valid storage object.",
                "PUB_DOC_01; PRIVATE_DOC_01; PENDING_DOC_01; DELETED_DOC_01.",
                [
                    "Open PUB_LIB_01 as Guest.",
                    "Select PUB_DOC_01 and click Download or open its public document route.",
                    "Call GET /api/public/documents/{documentId}/download.",
                    "Open the returned signed URL before expiry.",
                    "Wait until the signed URL expires and retry the same URL.",
                ],
                "The API returns a short-lived signed URL for PUB_DOC_01. The file downloads successfully before expiry and the expired URL is rejected by storage.",
                [
                    "Try the public endpoint with PRIVATE_DOC_01, PENDING_DOC_01 and DELETED_DOC_01; each must return 404/controlled denial.",
                    "Change only document.is_public to true while the parent library remains private; access must remain denied.",
                ],
                "FR-PUB-03; BR-DOC-07; FR-DOC-05",
            ),
            case(
                "UC-PUB-03",
                "Access as Guest",
                "Guest",
                "No authenticated session exists.",
                "Fresh browser profile; optionally an expired/stale token from a previous test.",
                [
                    "Open Login and select Access as Guest.",
                    "Inspect browser storage and network requests.",
                    "Navigate to Home, Discover, public Libraries, public Profile and public Document routes.",
                    "Attempt to navigate to AI Chat, Create Library, Workspace and Admin routes.",
                    "Attempt a protected mutating API without a valid JWT.",
                ],
                "Guest mode stores only restricted client state and no fake JWT. Public screens work; protected routes redirect or request sign-in; protected API returns 401/403.",
                [
                    "Verify no guest_signature_bypass token is created.",
                    "Verify Guest cannot create, update or delete Library/Workspace data.",
                    "Verify a stale invalid token is cleared without an infinite refresh loop.",
                ],
                "FR-PUB-04; BR-SEC-02; BR-SEC-03",
            ),
        ],
    ),
    (
        "B. Identity and Account",
        [
            case(
                "UC-AUTH-01",
                "Register or Sign In with Google",
                "Guest; Google Identity Provider",
                "Valid Google OAuth client and redirect configuration; email service available for first-time setup.",
                "New Google account; existing completed account; invalid/deleted OAuth client.",
                [
                    "Open Login/Register and select Google sign-in.",
                    "Authenticate with a new Google account.",
                    "Verify the backend validates the Google credential and starts the setup/OTP flow.",
                    "Repeat with an existing completed account.",
                    "Inspect the issued access/refresh session and destination page.",
                ],
                "New account enters verification/setup; existing account signs in and reaches the dashboard. Tokens are issued only after valid Google verification.",
                [
                    "Use an invalid Google token and expect AUTH-001.",
                    "Use a disabled account and expect access denial.",
                    "Test an invalid/deleted client ID and verify a controlled login error, not an application crash.",
                ],
                "FR-AUTH-01; BR-AUTH-05; BR-SEC-02",
            ),
            case(
                "UC-AUTH-02",
                "Verify Registration OTP and Complete Setup",
                "Guest; Email Service",
                "A pending account and active OTP exist.",
                "Valid OTP; wrong OTP; expired OTP; valid/duplicate username; valid/weak password.",
                [
                    "Submit the valid registration OTP within ten minutes.",
                    "Continue to Complete Profile.",
                    "Enter a unique username, password satisfying policy and profile data.",
                    "Submit setup and save a bio through /api/profile/me.",
                    "Sign in with the new credentials.",
                ],
                "OTP is consumed once, profile becomes active, username is unique, password is stored securely, bio persists, and the user can sign in.",
                [
                    "Wrong, expired or reused OTP is rejected.",
                    "Duplicate username and weak password are rejected.",
                    "Setup token older than 15 minutes is rejected.",
                ],
                "FR-AUTH-02; FR-AUTH-03; BR-AUTH-01; BR-AUTH-02; BR-AUTH-03; BR-AUTH-04",
            ),
            case(
                "UC-AUTH-03",
                "Log In with Credentials",
                "Guest / Registered User",
                "Active, disabled and invalid test accounts exist.",
                "Email or username; correct/incorrect passwords; remember-me on/off.",
                [
                    "Enter valid username/email and password.",
                    "Submit login with Remember Me disabled.",
                    "Confirm access token, refresh cookie and user profile are stored in the intended scopes.",
                    "Log out, repeat with Remember Me enabled and reopen the browser.",
                    "Verify authenticated routes and current profile.",
                ],
                "Active user signs in, receives an active session and reaches the dashboard. Remember Me behavior matches the selected option.",
                [
                    "Incorrect credentials return AUTH-004 without revealing which field is wrong.",
                    "Disabled user receives AUTH-005.",
                    "Repeated login attempts trigger the sensitive-auth rate limiter.",
                ],
                "FR-AUTH-04; BR-AUTH-04; BR-AUTH-06; BR-SEC-01",
            ),
            case(
                "UC-AUTH-04",
                "Refresh Session and Log Out",
                "Registered User / System Admin",
                "An authenticated session exists.",
                "Valid refresh cookie; expired/revoked refresh cookie; access token near expiry.",
                [
                    "Use the application until the access token approaches expiry.",
                    "Trigger a protected request and observe refresh behavior.",
                    "Verify the new access token retains the correct role and session_id.",
                    "Choose Log Out.",
                    "Attempt to reuse the old access and refresh credentials.",
                ],
                "A valid refresh session issues a new access token. Logout clears the refresh cookie/client session and old credentials cannot access protected APIs.",
                [
                    "Expired or revoked refresh token is rejected.",
                    "Login from another session invalidates the previous session when session_id changes.",
                    "Refresh failure redirects protected non-Guest users to Login without loops.",
                ],
                "FR-AUTH-05; BR-AUTH-04; BR-AUTH-06",
            ),
            case(
                "UC-AUTH-05",
                "Recover Password",
                "Guest; Email Service",
                "An active password account exists and email delivery is configured.",
                "Known/unknown email; valid/wrong/expired OTP; strong/weak new password.",
                [
                    "Submit Forgot Password for a registered email.",
                    "Receive and submit the reset OTP.",
                    "Use the issued reset token to set a new valid password.",
                    "Sign in with the new password.",
                    "Attempt sign-in with the old password.",
                ],
                "OTP and reset token are single-use and time-limited. New password works; old password fails.",
                [
                    "Wrong/expired/reused OTP is rejected.",
                    "Weak password and expired reset token are rejected.",
                    "Repeated requests trigger auth rate limiting without leaking account existence.",
                ],
                "FR-AUTH-06; BR-AUTH-01; BR-AUTH-03; BR-AUTH-04",
            ),
            case(
                "UC-AUTH-06",
                "Change Password",
                "Registered User",
                "Authenticated password-based account.",
                "Correct/incorrect current password; valid/invalid new passwords.",
                [
                    "Open account settings.",
                    "Enter the correct current password and a compliant new password.",
                    "Submit the change.",
                    "Log out and sign in with the new password.",
                    "Verify the old password no longer works.",
                ],
                "Password changes only after current-password verification and policy validation.",
                [
                    "Incorrect current password returns AUTH-006.",
                    "Weak new password returns AUTH-007.",
                    "Guest and unauthenticated requests are rejected.",
                ],
                "FR-AUTH-07; BR-AUTH-03; BR-SEC-02",
            ),
            case(
                "UC-AUTH-07",
                "Delete Account",
                "Registered User",
                "Authenticated user owns sample libraries/documents and may belong to workspaces.",
                "Password account and Google-only account; confirmation values DELETE/delete/blank.",
                [
                    "Open account deletion.",
                    "Enter the literal DELETE confirmation.",
                    "For a password account, provide the correct current password.",
                    "Submit and confirm the account/session is removed.",
                    "Attempt login and access to former private resources.",
                ],
                "The account is deleted according to data-retention rules, sessions are invalidated and former private resources are no longer accessible.",
                [
                    "Any confirmation other than DELETE is rejected with AUTH-008.",
                    "Wrong password is rejected.",
                    "Verify Workspace ownership constraints are handled without orphaning an active Workspace.",
                ],
                "FR-AUTH-08; BR-AUTH-08; BR-SEC-03",
            ),
        ],
    ),
    (
        "C. Profile and Discovery",
        [
            case(
                "UC-PROF-01",
                "Manage Own Profile and Avatar",
                "Registered User",
                "Authenticated active user.",
                "Valid profile fields; JPG/PNG/WEBP under 2 MB; invalid type; image over 2 MB.",
                [
                    "Call GET /api/profile/me and open the Profile page.",
                    "Update supported fields including full name and bio using PUT /api/profile/me.",
                    "Reload and verify persistence.",
                    "Upload a valid avatar.",
                    "Verify the new avatar appears in Profile and navigation.",
                ],
                "Only the authenticated profile is returned/updated. Valid avatar is stored and displayed consistently.",
                [
                    "Attempt to update another user ID through the payload; it must be ignored/rejected.",
                    "Reject unsupported MIME type and file over 2 MB.",
                    "Guest and disabled user cannot update a profile.",
                ],
                "FR-PROF-01; FR-PROF-02; BR-PROF-01; BR-SEC-03",
            ),
            case(
                "UC-PROF-02",
                "Search Users and View Public Profile",
                "Guest / Registered User",
                "Several active and disabled profiles exist with public-safe fields.",
                "Exact and partial username/name queries; unknown user ID.",
                [
                    "Search for an active user by username and display name.",
                    "Open the selected public profile.",
                    "Repeat from Guest mode.",
                    "Inspect all returned profile fields.",
                ],
                "Matching active profiles and only permitted public fields are returned. Public profile opens for Guest and Registered User.",
                [
                    "Disabled/private-sensitive data must not be exposed.",
                    "Unknown user returns controlled 404/empty results.",
                    "Verify password_hash, session_id and private account fields never appear.",
                ],
                "FR-PROF-03; BR-SEC-01; BR-SEC-03",
            ),
        ],
    ),
    (
        "D. Library and Documents",
        [
            case(
                "UC-LIB-01",
                "Create, Update or Delete Library",
                "Registered User",
                "Authenticated user has fewer than five libraries.",
                "Unique/duplicate/blank names; public/private; share_on_profile true/false.",
                [
                    "Create a Library with name, description, visibility and profile-sharing setting.",
                    "Open Library Settings and update its metadata.",
                    "Toggle public to private and inspect child document visibility.",
                    "Delete the Library after confirmation.",
                    "Verify dependent engagement records and document association follow deletion rules.",
                ],
                "Owner can manage only their Library. Visibility changes propagate to child document public visibility. Deletion occurs only after ownership validation.",
                [
                    "Blank or duplicate name is rejected.",
                    "Sixth active Library is rejected.",
                    "Guest, authenticated non-owner and System Admin owner-only API calls are rejected.",
                ],
                "FR-LIB-01; BR-LIB-01; BR-SEC-03",
            ),
            case(
                "UC-LIB-02",
                "List and Open Own Libraries",
                "Registered User",
                "User owns public and private Libraries; another user owns additional Libraries.",
                "OWNER_LIB_PUBLIC; OWNER_LIB_PRIVATE; OTHER_LIB_PUBLIC; OTHER_LIB_PRIVATE.",
                [
                    "Open My Libraries.",
                    "Call GET /api/documents/libraries.",
                    "Open each owned Library through GET /api/documents/libraries/{libraryId}.",
                    "Inspect counts, metadata and ownership labels.",
                ],
                "Only owned Libraries appear in My Libraries and owner endpoint. Public Library of another user is available only through the public read-only flow.",
                [
                    "Opening another user's private Library returns 404/403.",
                    "Authenticated non-owner sees no Settings, Save, Upload or Delete controls on public Library.",
                    "Direct update/delete requests by non-owner return 403.",
                ],
                "FR-LIB-02; BR-SEC-03",
            ),
            case(
                "UC-DOC-01",
                "Suggest Tags for Documents",
                "Registered User",
                "AI provider is available; supported test files contain readable text.",
                "PDF/DOCX/TXT with relevant content; empty file; unsupported file.",
                [
                    "Select one or more supported files before upload.",
                    "Request AI tag suggestions.",
                    "Inspect normalized, unique suggested tags.",
                    "Choose/edit tags and continue to upload.",
                ],
                "Suggestions are based on content, normalized and bounded. Failure produces a controlled retryable message.",
                [
                    "Reject unsupported file types and unreadable files.",
                    "Reject duplicate tags or tags beginning with a number.",
                    "Simulate AI quota/service failure and expect a controlled 503/defined error code.",
                ],
                "FR-DOC-01; FR-DOC-02; BR-DOC-01; BR-DOC-02; BR-DOC-04",
            ),
            case(
                "UC-DOC-02",
                "Upload Documents",
                "Registered User / Workspace Editor / Workspace Admin",
                "Target Library or Workspace exists and requester has required ownership/role.",
                "Valid PDF/DOCX/TXT; unsupported file; 50 MB boundary; duplicate; sensitive and unreadable content.",
                [
                    "Select valid files and target scope.",
                    "Submit upload with valid tags.",
                    "Verify file validation, quota and duplicate checks.",
                    "Inspect storage object and document row.",
                    "Wait for extraction, moderation, chunks, embeddings and final/pending status.",
                ],
                "Library upload is owner-only and inherits Library visibility. Workspace Editor/Admin upload returns quickly with PENDING then completes AI processing. Records, tags and chunks are linked correctly.",
                [
                    "Non-owner Library upload and Workspace Viewer/non-member upload return 403.",
                    "Unsupported, oversized, quota-exceeding or unreadable files are rejected.",
                    "Duplicate without explicit replacement confirmation returns DOC-002.",
                    "AI failure produces PENDING_RETRY rather than a frozen request.",
                ],
                "FR-DOC-03; FR-AI-01; BR-DOC-01..06; BR-DOC-08; BR-WS-03",
            ),
            case(
                "UC-DOC-03",
                "View or Download Document",
                "Registered User",
                "Private Library, public Library and Workspace documents exist in multiple statuses.",
                "Owner/non-owner; Workspace member/non-member; APPROVED/PENDING/FLAGGED/deleted.",
                [
                    "As owner, view and download a private Library document.",
                    "As authenticated non-owner, access an APPROVED public document.",
                    "As Workspace member, access an authorized Workspace document.",
                    "Inspect signed URL duration and bucket selection.",
                ],
                "Access is granted only by ownership, valid Workspace membership or approved public Library rules. Signed URLs are short-lived and point to the correct bucket.",
                [
                    "Non-member cannot access Workspace document by ID.",
                    "Non-owner cannot access private Library document even if document.is_public is stale true.",
                    "Deleted document and unauthorized AI access are denied.",
                ],
                "FR-DOC-04; FR-DOC-05; BR-DOC-07; BR-WS-05; BR-SEC-03",
            ),
            case(
                "UC-DOC-04",
                "Delete or Replace Document",
                "Registered User / Workspace Admin",
                "Owned and Workspace documents exist; duplicate detection data prepared.",
                "Original uploader; different Workspace Editor; Workspace Admin; replacement confirmation IDs.",
                [
                    "Soft-delete an owned Library document.",
                    "Upload a duplicate without confirmation and inspect conflict information.",
                    "Confirm replacement with the returned document ID.",
                    "As Workspace Admin, replace/remove an authorized Workspace document.",
                    "Verify old/new rows, storage state and visibility.",
                ],
                "Normal deletion sets deleted_at. Replacement occurs only for confirmed targets and within the same Library/Workspace scope.",
                [
                    "Non-owner cannot delete Library document.",
                    "Workspace Editor cannot replace another uploader's document unless SRS-authorized role permits it.",
                    "A replacement ID from another scope is rejected.",
                ],
                "FR-DOC-06; BR-DOC-03; BR-SEC-03",
            ),
        ],
    ),
    (
        "E. AI Study",
        [
            case(
                "UC-AI-01",
                "Process and Moderate Document",
                "Google Gemini AI / System Flow",
                "A document upload has created a document row and storage object.",
                "Readable safe file; sensitive file; file under 20 readable characters; simulated embedding failure.",
                [
                    "Trigger processing for a readable safe document.",
                    "Verify extraction, moderation, chunk splitting and embeddings.",
                    "Repeat with sensitive content.",
                    "Repeat with unreadable content.",
                    "Simulate AI/embedding provider failure.",
                ],
                "Safe content produces chunks and appropriate status. Sensitive content enters review status. Unreadable content is REJECTED. Provider failure becomes PENDING_RETRY.",
                [
                    "No document may be APPROVED for AI use without chunks.",
                    "Verify sensitive file is stored in the review bucket.",
                    "Retry processing does not create duplicate active chunks.",
                ],
                "FR-AI-01; BR-DOC-05; BR-AI-01",
            ),
            case(
                "UC-AI-02",
                "Ask Document-Grounded Question",
                "Registered User",
                "Accessible APPROVED document has chunks/embeddings; daily quota not exhausted.",
                "Relevant question; unrelated question; private/Workspace unauthorized document; 50/51 quota boundary.",
                [
                    "Select an approved accessible document.",
                    "Submit a relevant question.",
                    "Inspect the grounded answer and source chunk hints.",
                    "Submit an unrelated question.",
                    "Inspect daily usage increment.",
                ],
                "Answer is generated from retrieved document chunks, includes source hints where implemented and increments chat usage once.",
                [
                    "PENDING/REJECTED/missing-chunk document returns AI-001/AI-002.",
                    "Unauthorized private or Workspace document is denied.",
                    "The 51st daily request returns AI-003/429.",
                ],
                "FR-AI-02; FR-AI-05; BR-AI-01; BR-AI-02",
            ),
            case(
                "UC-AI-03",
                "Generate Flashcards",
                "Registered User",
                "Accessible APPROVED document has chunks and user has remaining daily allowance.",
                "Documents with sufficient/insufficient content; daily generated-card count 0, 2 and 3.",
                [
                    "Select an approved document.",
                    "Request flashcard generation.",
                    "Verify generated question/answer cards are saved and linked to user/document/workspace where applicable.",
                    "Open Flashcards and study the generated set.",
                ],
                "Up to the remaining daily allowance is generated and stored from document chunks.",
                [
                    "Non-approved, unauthorized or missing-chunk document is rejected.",
                    "Request after three cards in the Vietnam day returns AI-004.",
                    "AI malformed output is handled without storing invalid cards.",
                ],
                "FR-AI-03; BR-AI-01; BR-AI-03",
            ),
            case(
                "UC-AI-04",
                "View AI Usage Summary",
                "Registered User",
                "AI usage records exist for today and prior days.",
                "0, partial and full daily chat usage; token counts.",
                [
                    "Open Home/AI usage summary.",
                    "Call GET /api/ai/summary.",
                    "Compare chats used, remaining and token totals with database records.",
                    "Cross the Vietnam-day boundary and refresh.",
                ],
                "Summary reports the current user's daily limit, used/remaining chats and non-negative token count; a new day resets daily usage.",
                [
                    "User cannot retrieve another user's summary.",
                    "Guest does not invoke protected AI usage APIs.",
                    "Missing usage row returns zero values rather than an error.",
                ],
                "FR-AI-04; BR-AI-02",
            ),
        ],
    ),
    (
        "F. Workspace Collaboration",
        [
            case(
                "UC-WS-01",
                "Create, View, Update or Delete Workspace",
                "Registered User / Workspace Admin",
                "Registered user owns fewer than three active Workspaces.",
                "Valid/blank names; owner with 0, 2 and 3 active Workspaces.",
                [
                    "Create a Workspace with name and description.",
                    "Verify workspaces.created_by and Admin membership.",
                    "Open Workspace detail.",
                    "As Admin, update name/description.",
                    "Soft-delete the Workspace and verify it disappears from active lists.",
                ],
                "Workspace and Admin membership are created atomically. Only members can view; only Admin can update/delete.",
                [
                    "Guest and unauthenticated creation are rejected.",
                    "Fourth owned active Workspace is rejected.",
                    "Viewer, Editor and non-member cannot rename/delete.",
                ],
                "FR-WS-01; BR-WS-01; BR-WS-02; BR-ADM-04",
            ),
            case(
                "UC-WS-02",
                "Manage Workspace Members",
                "Workspace Admin",
                "Workspace has Admin, Editor, Viewer and eligible non-member users.",
                "Valid/duplicate invite; roles Editor/Viewer/invalid; transfer target member.",
                [
                    "Search eligible users and invite/add an Editor and Viewer.",
                    "Accept an invitation as the target user.",
                    "Change Viewer to Editor and back.",
                    "Transfer ownership to an existing member.",
                    "Remove a non-owner member.",
                ],
                "Membership and role changes persist. Ownership transfer atomically updates target role, former owner role and workspaces.created_by.",
                [
                    "Non-Admin cannot manage members.",
                    "Duplicate member/invite and invalid role are rejected.",
                    "Invitation to a deleted Workspace cannot be accepted.",
                    "Ownership transfer failure rolls back all ownership/role changes.",
                ],
                "FR-WS-02; FR-WS-03; BR-WS-02; BR-WS-06",
            ),
            case(
                "UC-WS-03",
                "Send and Read Workspace Messages",
                "Workspace Member",
                "Active Workspace with Admin, Editor and Viewer members.",
                "Normal text; blank/oversized text if a limit exists; non-member token.",
                [
                    "Open Workspace Messages as Viewer.",
                    "Read existing messages.",
                    "Post a new message.",
                    "Refresh as Editor/Admin and verify sender, content and time.",
                ],
                "All active Workspace members can read and post messages; messages remain scoped to the Workspace.",
                [
                    "Non-member cannot list or post messages.",
                    "Blank/invalid message is rejected.",
                    "Soft-deleted Workspace messages are inaccessible.",
                ],
                "FR-WS-04; BR-WS-05",
            ),
            case(
                "UC-WS-04",
                "Manage Workspace Discussions",
                "Workspace Member / Workspace Editor / Workspace Admin",
                "Workspace contains Admin, two Editors and one Viewer.",
                "Topic title/content/status/priority; comments and solutions.",
                [
                    "As Editor A, create a discussion Topic.",
                    "As Editor B, update the Topic title/status/priority.",
                    "As Viewer, read the Topic and submit a solution/comment.",
                    "As Editor B, delete the Topic.",
                    "Verify soft deletion and list behavior.",
                ],
                "All members read and submit solutions. Editor and Admin can create/update/delete any Topic as required by SRS.",
                [
                    "Viewer cannot create/update/delete Topic.",
                    "Non-member cannot read discussion content.",
                    "Blank title and invalid Topic ID return controlled errors.",
                ],
                "FR-WS-06; BR-WS-04; BR-WS-05",
            ),
            case(
                "UC-WS-05",
                "Manage Discussion Subtasks and Attachments",
                "Workspace Member / Workspace Editor / Workspace Admin",
                "An active discussion Topic exists.",
                "Subtask title/done state; supported attachment; solution/comment attachment.",
                [
                    "As Editor, add a Subtask to the Topic.",
                    "Update its title/completion state and then delete it.",
                    "Attach a file to Topic/solution/comment through the Workspace upload path.",
                    "Refresh and verify attachment metadata and download access.",
                ],
                "Authorized Editor/Admin manages Subtasks. Members participate through solution/comment attachments according to endpoint authorization. Attachments appear promptly and process in background.",
                [
                    "Viewer cannot perform Editor-only Subtask management.",
                    "Attachment with wrong workspaceId or by non-member is rejected.",
                    "Unsupported/oversized attachment is rejected without freezing the UI.",
                ],
                "FR-WS-07; BR-WS-04; BR-DOC-08",
            ),
            case(
                "UC-WS-06",
                "Upload and Review Workspace Documents",
                "Workspace Editor / Workspace Admin",
                "Workspace storage is below 50 MB; Editor and Admin memberships exist.",
                "Safe, sensitive and invalid files; storage boundary; APPROVE/REJECT decisions.",
                [
                    "As Editor, upload a valid Workspace document.",
                    "Verify immediate PENDING response and background extraction/moderation/chunking.",
                    "As Workspace Admin, inspect the document review queue.",
                    "Approve a safe document and verify movement to approved storage/status.",
                    "Reject another document with a reason and verify review storage/status.",
                ],
                "Editor/Admin upload; Viewer cannot. Admin review records reviewer, timestamp and reason. Approved document becomes usable for view/AI when chunks exist.",
                [
                    "Invalid review decision is rejected.",
                    "Non-Admin cannot review.",
                    "Storage move failure must not report a successful status update.",
                    "Workspace over 50 MB rejects additional upload.",
                ],
                "FR-WS-08; BR-DOC-08; BR-WS-02; BR-WS-03",
            ),
            case(
                "UC-WS-07",
                "View Workspace Flashcards and Notifications",
                "Workspace Member / Registered User",
                "Workspace contains flashcards and membership/activity notifications.",
                "Admin, Editor, Viewer and non-member accounts.",
                [
                    "Open Workspace flashcards as Viewer, Editor and Admin.",
                    "Study the available cards.",
                    "Trigger invitation, role-change and document-review events.",
                    "Open Notifications and mark items as read where supported.",
                ],
                "Members can view Workspace flashcards. Relevant authenticated users receive correctly scoped notifications.",
                [
                    "Non-member cannot access Workspace flashcards.",
                    "User cannot read another user's private notifications.",
                    "Guest does not call protected notification APIs.",
                ],
                "FR-WS-09; BR-WS-05",
            ),
        ],
    ),
    (
        "G. System Administration",
        [
            case(
                "UC-ADM-01",
                "View Admin Dashboard",
                "System Admin",
                "System contains representative users, documents, usage and moderation data.",
                "SYSTEM_ADMIN, USER and Guest sessions.",
                [
                    "Sign in as System Admin and open /admin/dashboard.",
                    "Inspect aggregate cards, charts and recent operational data.",
                    "Compare totals with API/database fixtures.",
                    "Refresh and test empty-data state.",
                ],
                "Dashboard loads accurate aggregate user, document, moderation, storage and AI metrics for System Admin only.",
                [
                    "USER, Workspace Admin, Guest and unauthenticated access receive redirect/403.",
                    "API failure shows a controlled error state.",
                    "No private secrets or password data appear in metrics.",
                ],
                "FR-ADM-01; BR-ADM-01",
            ),
            case(
                "UC-ADM-02",
                "Review Moderated Documents",
                "System Admin",
                "Moderation queue contains FLAGGED, REJECTED and PENDING_RETRY documents.",
                "APPROVE/KEEP_REJECTED; valid/blank reason; storage transfer failure.",
                [
                    "Open moderation queue and filter/search items.",
                    "Inspect document context and moderation reason.",
                    "Approve one flagged item with a reason.",
                    "Keep another item rejected with a reason.",
                    "Verify status, storage, chunks/tags and audit log.",
                ],
                "Valid decision updates document lifecycle consistently. Approve transfers file before status success; KEEP_REJECTED removes required data and records reason/audit.",
                [
                    "Invalid decision or blank reason is rejected.",
                    "Storage failure does not mark approval successful.",
                    "Non-System Admin receives 403.",
                ],
                "FR-ADM-02; BR-ADM-01; BR-ADM-02",
            ),
            case(
                "UC-ADM-03",
                "Change User Status",
                "System Admin",
                "Active and disabled users exist; current admin identified.",
                "Target ACTIVE/DISABLED; self-target; last active admin.",
                [
                    "Open User Management.",
                    "Disable an active normal user.",
                    "Attempt login/protected request as that user.",
                    "Reactivate the user and sign in again.",
                    "Inspect activity log.",
                ],
                "Status change persists, disabled account loses protected access, reactivated account can authenticate, and high-impact action is logged.",
                [
                    "Admin cannot disable their own account.",
                    "Non-admin cannot change status.",
                    "System must not leave zero active System Admin accounts.",
                ],
                "FR-ADM-03; BR-ADM-01; BR-ADM-03",
            ),
            case(
                "UC-ADM-04",
                "Change System Role",
                "System Admin",
                "At least two active System Admins and normal users exist.",
                "USER to SYSTEM_ADMIN; SYSTEM_ADMIN to USER; self; final active admin.",
                [
                    "Promote a USER to SYSTEM_ADMIN.",
                    "Sign in as the promoted account and verify admin access.",
                    "Demote a non-final SYSTEM_ADMIN to USER.",
                    "Verify admin access is removed.",
                    "Inspect audit logs.",
                ],
                "Role changes persist and frontend/backend authorization changes immediately or at the next validated session.",
                [
                    "Admin cannot change their own system role.",
                    "Final active System Admin cannot be demoted.",
                    "Unsupported role value and non-admin request are rejected.",
                ],
                "FR-ADM-04; BR-ADM-01; BR-ADM-03",
            ),
            case(
                "UC-ADM-05",
                "View Activity Logs",
                "System Admin",
                "Audit events exist across authentication, moderation, user, Workspace and issue workflows.",
                "Filters by actor/action/entity/risk/date; pagination boundaries.",
                [
                    "Open Activity Logs.",
                    "Filter by action type, risk level, actor and date.",
                    "Inspect old/new data, IP, device/user-agent and details.",
                    "Navigate across pagination and clear filters.",
                ],
                "Only matching events are returned with stable pagination and complete permitted audit metadata.",
                [
                    "Non-System Admin cannot access logs.",
                    "Invalid filter input is controlled and does not alter queries unsafely.",
                    "Sensitive secrets/passwords are not stored in log payloads.",
                ],
                "FR-ADM-05; BR-ADM-01",
            ),
            case(
                "UC-ADM-06",
                "Monitor System Usage",
                "System Admin",
                "Upload/download quota and AI usage data exist for multiple users/dates.",
                "User/date filters; zero usage; high usage; pagination.",
                [
                    "Open Admin Usage.",
                    "Filter usage by user and date.",
                    "Compare upload/download bytes, chat counts and tokens with fixtures.",
                    "Verify quota/limit calculations and empty state.",
                ],
                "Usage data is accurate, filterable and paginated as implemented; byte/token/count values remain non-negative.",
                [
                    "Non-System Admin receives 403.",
                    "Invalid dates/identifiers return controlled validation.",
                    "One user's usage must not be attributed to another.",
                ],
                "FR-ADM-06; BR-ADM-01",
            ),
            case(
                "UC-ADM-07",
                "Manage Deleted Workspaces",
                "System Admin",
                "Soft-deleted and active Workspaces exist; owner-limit and storage fixtures prepared.",
                "Restore eligible/ineligible Workspace; exact/wrong confirmation; workspace-only and Library-linked documents.",
                [
                    "Open Deleted Workspaces and review purge preview.",
                    "Restore an eligible soft-deleted Workspace.",
                    "Soft-delete another Workspace and enter its exact name for permanent purge.",
                    "Verify workspace-only files are removed before database purge.",
                    "Verify Library-linked documents are preserved and detached.",
                ],
                "Only soft-deleted Workspaces can be restored/purged. Restore respects the three-owned limit. Confirmed purge follows BR-ADM-06 and is logged.",
                [
                    "Wrong confirmation and active Workspace purge are rejected.",
                    "Restore is rejected when owner already has three active Workspaces.",
                    "Storage cleanup failure leaves database records for retry and does not report success.",
                    "Non-System Admin receives 403.",
                ],
                "FR-ADM-07; FR-ADM-08; BR-ADM-04..06",
            ),
            case(
                "UC-ADM-08",
                "Manage Issue Reports",
                "System Admin",
                "Issue queue contains OPEN, IN_PROGRESS, RESOLVED and DISMISSED reports.",
                "Status, priority and search filters; valid/invalid admin response.",
                [
                    "Open Admin Issue Reports.",
                    "Search/filter by status, category and priority.",
                    "Open a report and inspect reporter/context.",
                    "Update priority, status and administrator response.",
                    "Set RESOLVED and inspect resolved_at/handler.",
                ],
                "Updates persist, reporter can see the response/status, and RESOLVED sets resolution metadata.",
                [
                    "Unsupported status/priority is rejected.",
                    "Non-System Admin cannot access the queue.",
                    "Admin cannot alter reporter ownership through the payload.",
                ],
                "FR-ADM-09; FR-ADM-10; BR-SUP-01; BR-SUP-03",
            ),
        ],
    ),
    (
        "H. Support",
        [
            case(
                "UC-SUP-01",
                "Report and Track System Issue",
                "Registered User",
                "Authenticated user and System Admin exist.",
                "Categories BUG/ACCOUNT/WORKSPACE/DOCUMENT/AI/OTHER; boundary-length fields.",
                [
                    "Open Report Issue.",
                    "Enter supported category, title, description, optional reproduction steps and page path.",
                    "Submit and verify OPEN/NORMAL defaults.",
                    "Open My Reports and inspect the submitted item.",
                    "After Admin update, refresh and verify status/response.",
                ],
                "Valid report is stored with the authenticated reporter, timestamps, OPEN status and NORMAL priority. User sees only their reports and later admin response.",
                [
                    "Reject unsupported category and title/description outside limits.",
                    "Reject optional steps over 3000 and page path over 500 characters.",
                    "User A cannot list or open User B's report.",
                    "Guest cannot submit an issue report.",
                ],
                "FR-SUP-01; FR-SUP-02; BR-SUP-01..03",
            ),
        ],
    ),
]


ROLE_MATRIX = [
    ("Landing/Auth/Recovery", "Yes", "Yes", "Yes", "Yes", "Yes", "Yes"),
    ("Browse public libraries/documents", "Yes", "Yes", "Yes", "Yes", "Yes", "Yes"),
    ("Own Library/Profile management", "No", "Yes", "Yes", "Yes", "Yes", "Inherited user scope"),
    ("Workspace detail/messages", "No", "No", "Yes", "Yes", "Yes", "No"),
    ("Workspace Topic create/edit/delete", "No", "No", "No", "Yes", "Yes", "No"),
    ("Workspace document upload", "No", "No", "No", "Yes", "Yes", "No"),
    ("Workspace member/document review", "No", "No", "No", "No", "Yes", "No"),
    ("Admin governance", "No", "No", "No", "No", "No", "Yes"),
    ("Report system issue", "No", "Yes", "Yes", "Yes", "Yes", "Yes"),
]


TEST_ACCOUNTS = [
    ("GUEST-01", "Guest", "No account/token", "Public browsing and authorization denial"),
    ("USER-01", "Registered User", "ACTIVE; owns public/private Libraries", "Profile, Library, Document, AI, Support"),
    ("USER-02", "Registered User", "ACTIVE; non-owner", "Cross-user authorization"),
    ("WS-VIEW-01", "Workspace Contributor (Viewer)", "Member of WS-01", "Read/message/solution; denied management"),
    ("WS-EDIT-01", "Workspace Editor", "Editor of WS-01", "Topic/Subtask/document upload"),
    ("WS-EDIT-02", "Workspace Editor", "Second Editor of WS-01", "Cross-editor Topic management"),
    ("WS-ADMIN-01", "Workspace Admin", "Creator/owner of WS-01", "Membership, review, ownership transfer"),
    ("SYS-ADMIN-01", "System Admin", "ACTIVE", "All administration use cases"),
    ("SYS-ADMIN-02", "System Admin", "ACTIVE backup admin", "Final-admin protection"),
    ("DISABLED-01", "Registered User", "DISABLED", "Authentication/status denial"),
]


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Bullet 2", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("AI STUDYHUB  |  COMPREHENSIVE TEST PLAN")
    set_run_font(run, size=9, bold=True, color=MID_GRAY)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("SRS v3.3  |  Test execution copy")
    set_run_font(run, size=9, color=MID_GRAY)


def add_cover(doc):
    for _ in range(5):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker = p.add_run("SOFTWARE TEST PLAN")
    set_run_font(kicker, size=11, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    title = p.add_run("AI StudyHub")
    set_run_font(title, size=30, bold=True, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    subtitle = p.add_run("Comprehensive Use Case & Role-Based Test Plan")
    set_run_font(subtitle, size=15, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Based on AI StudyHub SRS v3.3")
    set_run_font(run, size=11, bold=True, color=MID_GRAY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Prepared: {date.today().strftime('%d %B %Y')}")
    set_run_font(run, size=10.5, color=MID_GRAY)

    for _ in range(8):
        doc.add_paragraph()

    add_callout(
        doc,
        "Coverage",
        "38 use cases; Guest, Registered User, Workspace Contributor, Editor, Admin and System Admin; positive, negative, authorization, security and lifecycle checks.",
        LIGHT_BLUE,
    )
    doc.add_page_break()


def add_document_control(doc):
    add_heading(doc, "1. Document Control", 1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2700, 6660])
    rows = [
        ("Document", "AI StudyHub Comprehensive Test Plan"),
        ("SRS baseline", "AI_StudyHub_SRS_v3.3_Visual_Paradigm_Use_Cases.docx"),
        ("SRS source", SRS_PATH),
        ("Scope", "Functional, authorization, negative, security and integration testing"),
        ("Status", "Ready for test execution after environment provisioning"),
        ("Pass definition", "Observed result matches every expected result and no unauthorized side effect occurs"),
    ]
    for index, (label, value) in enumerate(rows):
        cells = table.rows[0].cells if index == 0 else table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        set_cell_shading(cells[0], LIGHT_GRAY)
        set_run_font(cells[0].paragraphs[0].runs[0], size=10, bold=True, color=DARK_BLUE)
        set_run_font(cells[1].paragraphs[0].runs[0], size=10, color=INK)


def add_strategy(doc):
    add_heading(doc, "2. Test Objectives and Strategy", 1)
    add_callout(
        doc,
        "Primary objective",
        "Demonstrate that every SRS use case works for its intended actor and is rejected for unauthorized actors, while data, storage and AI lifecycle state remain consistent.",
    )
    add_heading(doc, "2.1 In scope", 2)
    for item in [
        "All 38 use cases in the SRS use-case catalog.",
        "Frontend route behavior and backend API enforcement.",
        "Library ownership, Workspace membership/roles and System Admin governance.",
        "Document storage, moderation, signed URLs, AI chunks/embeddings and status transitions.",
        "Happy path, validation boundaries, authorization bypass attempts and external-service failures.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2.2 Test levels", 2)
    for item in [
        "API integration: status code, response schema, authorization and database side effects.",
        "UI end-to-end: navigation, visibility of controls, feedback, loading and recovery.",
        "Data integrity: relational rows, soft-delete fields, role/owner consistency and audit logs.",
        "Storage/AI integration: bucket movement, signed URLs, retries and provider failure handling.",
        "Regression: backend automated suite, frontend lint/build and focused security tests.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2.3 Entry criteria", 2)
    for item in [
        "Required migrations are applied, including workspace ownership transaction and Library visibility reconciliation.",
        "OAuth, SMTP, Supabase Storage and Gemini test configuration is available or stubbed.",
        "Seed accounts and datasets in Sections 4 and 5 are prepared.",
        "Frontend and backend versions under test are recorded.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "2.4 Exit criteria", 2)
    for item in [
        "100% of Must/High use cases executed; all authorization checks executed.",
        "No open Critical/High defect affecting authentication, ownership, membership, privacy, purge or data loss.",
        "At least 95% overall test cases passed; remaining failures accepted with documented risk.",
        "Backend tests, frontend lint and production build pass for the release candidate.",
    ]:
        add_bullet(doc, item)


def add_role_matrix(doc):
    add_heading(doc, "3. Actor and Permission Coverage", 1)
    add_callout(
        doc,
        "Role mapping",
        "Workspace Contributor in the SRS maps to backend/database role Viewer. Workspace Admin inherits Editor capabilities. System Admin is a platform role and does not become owner of personal Libraries.",
        AMBER,
    )
    headers = [
        "Capability",
        "Guest",
        "Registered",
        "WS Contributor",
        "WS Editor",
        "WS Admin",
        "System Admin",
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    widths = [2700, 900, 1050, 1200, 1050, 1050, 1410]
    set_table_geometry(table, widths)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.runs[0], size=8.5, bold=True, color=DARK_BLUE)
    set_repeat_table_header(table.rows[0])
    for row_data in ROLE_MATRIX:
        row = table.add_row()
        for index, value in enumerate(row_data):
            row.cells[index].text = value
            p = row.cells[index].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(p.runs[0], size=8.5, color=INK)


def add_test_data(doc):
    add_heading(doc, "4. Test Environment and Accounts", 1)
    add_heading(doc, "4.1 Recommended environment", 2)
    for item in [
        "Isolated Supabase project or schema with disposable Storage buckets.",
        "Backend configured with strict CORS allow-list and production-like rate limits.",
        "Frontend production build plus browser DevTools for storage/network inspection.",
        "Test Google OAuth client, SMTP inbox and Gemini quota-safe project or deterministic mocks.",
        "Clock control or database fixtures for OTP/token/quota expiry and Vietnam-day boundaries.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "4.2 Test accounts", 2)
    headers = ["ID", "Role", "State", "Primary purpose"]
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    set_table_geometry(table, [1500, 2100, 2300, 3460])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        set_cell_shading(cell, LIGHT_BLUE)
        set_run_font(cell.paragraphs[0].runs[0], size=9, bold=True, color=DARK_BLUE)
    set_repeat_table_header(table.rows[0])
    for row_data in TEST_ACCOUNTS:
        row = table.add_row()
        for index, value in enumerate(row_data):
            row.cells[index].text = value
            set_run_font(row.cells[index].paragraphs[0].runs[0], size=9, color=INK)

    add_heading(doc, "4.3 Core data fixtures", 2)
    for item in [
        "Libraries: public/private, with child documents in APPROVED, PENDING, FLAGGED, REJECTED and deleted states.",
        "Workspace WS-01: Admin, two Editors, Viewer, non-member, messages, Topics, Subtasks, solutions and flashcards.",
        "Documents: PDF/DOCX/TXT; unsupported extension; 0/19/20+ readable characters; sensitive content; duplicate names/content; 50 MB boundaries.",
        "Administration: two active System Admins, disabled user, moderation queue, usage records, logs, soft-deleted Workspaces and issue reports.",
    ]:
        add_bullet(doc, item)


def add_execution_rules(doc):
    add_heading(doc, "5. Test Execution Rules", 1)
    add_heading(doc, "5.1 Evidence required", 2)
    for item in [
        "Record build/commit, environment, tester, date, browser and account ID.",
        "Capture request/response, UI screenshot and relevant database/storage state for failures.",
        "For authorization tests, verify both response denial and absence of unauthorized data mutation.",
        "For async AI tests, record status timeline and final state rather than only upload response.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "5.2 Result classification", 2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_geometry(table, [1500, 2500, 5360])
    headers = ["Result", "Meaning", "Required action"]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
        set_cell_shading(table.rows[0].cells[index], LIGHT_BLUE)
        set_run_font(table.rows[0].cells[index].paragraphs[0].runs[0], size=9.5, bold=True, color=DARK_BLUE)
    rows = [
        ("PASS", "All expected results met", "Attach evidence for critical authorization/lifecycle tests."),
        ("FAIL", "One or more expected results not met", "Create defect with severity, reproduction and evidence."),
        ("BLOCKED", "Environment/dependency prevents execution", "Record blocker owner and retest condition."),
        ("NOT RUN", "Not yet executed", "Keep scheduled before exit decision."),
    ]
    for result, meaning, action in rows:
        row = table.add_row()
        for index, value in enumerate((result, meaning, action)):
            row.cells[index].text = value
            set_run_font(row.cells[index].paragraphs[0].runs[0], size=9.5, color=INK)
        set_cell_shading(row.cells[0], GREEN if result == "PASS" else AMBER if result in ("BLOCKED", "NOT RUN") else RED)


def add_test_cases(doc):
    add_heading(doc, "6. Detailed Use Case Test Procedures", 1)
    add_callout(
        doc,
        "Execution rule",
        "Run each main procedure with the stated actor, then execute every negative/authorization check using the named alternate roles. A denial passes only when no database or storage side effect occurs.",
    )

    sequence = 0
    for group_title, cases in GROUPS:
        add_heading(doc, group_title, 2)
        for test_case in cases:
            sequence += 1
            add_heading(
                doc,
                f"TC-{sequence:03d} | {test_case['id']} - {test_case['title']}",
                3,
            )
            add_label_value(doc, "Primary actor", test_case["actor"])
            add_label_value(doc, "Priority", test_case["priority"])
            add_label_value(doc, "SRS traceability", test_case["rules"])
            add_label_value(doc, "Preconditions", test_case["preconditions"])
            add_label_value(doc, "Test data", test_case["data"])

            p = doc.add_paragraph()
            p.paragraph_format.keep_with_next = True
            run = p.add_run("Test steps")
            set_run_font(run, size=10.5, bold=True, color=DARK_BLUE)
            for step in test_case["steps"]:
                add_numbered(doc, step)

            add_callout(doc, "Expected result", test_case["expected"], GREEN)

            p = doc.add_paragraph()
            p.paragraph_format.keep_with_next = True
            run = p.add_run("Negative and authorization checks")
            set_run_font(run, size=10.5, bold=True, color=DARK_BLUE)
            for check in test_case["negative"]:
                add_bullet(doc, check)

            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(10)
            run = p.add_run("Execution result:  [ ] PASS   [ ] FAIL   [ ] BLOCKED   Defect ID: __________")
            set_run_font(run, size=9.5, italic=True, color=MID_GRAY)


def add_cross_cutting(doc):
    add_heading(doc, "7. Cross-Cutting Security and Reliability Tests", 1)
    cross_tests = [
        (
            "Authentication and session",
            "Protected APIs reject missing, expired, malformed and fake Guest tokens; disabled/session-invalidated users lose access; refresh/logout cannot be replayed.",
        ),
        (
            "Horizontal authorization",
            "Swap Library, Document, Workspace, Topic, report and user IDs between accounts; every owner/member/reporter boundary remains enforced by backend.",
        ),
        (
            "Role escalation",
            "Modify role fields/client storage and call Admin/Workspace Admin endpoints directly; server-side role checks remain authoritative.",
        ),
        (
            "CORS and rate limit",
            "Allowed frontend origins succeed; unknown origins fail; general and sensitive-auth limits trigger at configured thresholds behind the production proxy.",
        ),
        (
            "Data and storage consistency",
            "Inject failures between storage and database operations for upload, moderation, replacement and purge; verify controlled retry state and no falsely reported success.",
        ),
        (
            "AI/provider failure",
            "Simulate Gemini timeout, quota exhaustion, malformed response and embedding failure; requests return controlled errors and documents become PENDING_RETRY where specified.",
        ),
        (
            "Concurrency",
            "Run duplicate uploads, simultaneous ownership transfer/member changes and parallel moderation decisions; verify uniqueness and transaction protection.",
        ),
        (
            "Privacy",
            "Verify public/profile/admin responses never expose password hashes, refresh tokens, service credentials or unrelated private user data.",
        ),
    ]
    for title, detail in cross_tests:
        add_label_value(doc, title, detail)


def add_traceability(doc):
    add_heading(doc, "8. Use Case Traceability Matrix", 1)
    headers = ["Test ID", "Use case", "Actor", "Group", "Priority"]
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    set_table_geometry(table, [1000, 1500, 3100, 2360, 1400])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        set_cell_shading(cell, LIGHT_BLUE)
        set_run_font(cell.paragraphs[0].runs[0], size=9, bold=True, color=DARK_BLUE)
    set_repeat_table_header(table.rows[0])

    sequence = 0
    for group_title, cases in GROUPS:
        for test_case in cases:
            sequence += 1
            row = table.add_row()
            values = (
                f"TC-{sequence:03d}",
                test_case["id"],
                test_case["actor"],
                group_title.split(". ", 1)[-1],
                test_case["priority"],
            )
            for index, value in enumerate(values):
                row.cells[index].text = value
                set_run_font(row.cells[index].paragraphs[0].runs[0], size=8.5, color=INK)


def add_release_checklist(doc):
    add_heading(doc, "9. Release Regression Checklist", 1)
    checks = [
        "Backend automated tests pass.",
        "Frontend lint passes.",
        "Frontend production build passes with no JavaScript chunk above 500 KB.",
        "Database migrations are applied and RPC permissions verified.",
        "Guest creates no fake token and completes all public-use cases.",
        "Library owner/non-owner matrix passes for UI and API.",
        "Workspace Viewer/Editor/Admin matrix passes for every endpoint.",
        "Document public/private/Workspace authorization passes for view, download, AI and flashcards.",
        "Ownership transfer rollback and permanent purge failure injection pass.",
        "System Admin final-admin protection and audit logging pass.",
        "No unresolved Critical/High defect remains.",
    ]
    for check in checks:
        add_bullet(doc, f"[ ] {check}")


def build():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    configure_styles(doc)
    add_header_footer(doc)
    add_cover(doc)
    add_document_control(doc)
    add_strategy(doc)
    add_role_matrix(doc)
    add_test_data(doc)
    add_execution_rules(doc)
    add_test_cases(doc)
    add_cross_cutting(doc)
    add_traceability(doc)
    add_release_checklist(doc)

    core = doc.core_properties
    core.title = "AI StudyHub Comprehensive Test Plan"
    core.subject = "SRS v3.3 Use Case and Role-Based Test Plan"
    core.author = "AI StudyHub QA"
    core.keywords = "AI StudyHub, Test Plan, SRS, Use Case, RBAC"
    core.comments = "Generated from AI StudyHub SRS v3.3."

    doc.save(OUTPUT)
    print(OUTPUT)
    print(f"Use cases: {sum(len(cases) for _, cases in GROUPS)}")
    print(f"Paragraphs: {len(doc.paragraphs)}")
    print(f"Tables: {len(doc.tables)}")


if __name__ == "__main__":
    build()
