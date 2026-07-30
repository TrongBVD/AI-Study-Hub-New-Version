from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "IMPORT_LIBRARY_CODE_CHANGES_ANALYSIS.txt"
OUTPUT = ROOT / "IMPORT_LIBRARY_CODE_CHANGES_ANALYSIS.docx"

NAVY = "2D4059"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "687386"
LIGHT_BLUE = "E8EEF5"
CODE_FILL = "F3F5F7"
CODE_BORDER = "D7DDE5"
WHITE = "FFFFFF"


def set_run_font(run, name="Calibri", size=11, color="000000", bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_shading(paragraph, fill, border=None):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    if border:
        borders = OxmlElement("w:pBdr")
        for side in ("top", "left", "bottom", "right"):
            edge = OxmlElement(f"w:{side}")
            edge.set(qn("w:val"), "single")
            edge.set(qn("w:sz"), "4")
            edge.set(qn("w:space"), "4")
            edge.set(qn("w:color"), border)
            borders.append(edge)
        p_pr.append(borders)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("AI STUDY HUB  |  TECHNICAL REFERENCE")
    set_run_font(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run("Import Library code analysis  |  ")
    set_run_font(run, size=9, color=MUTED)
    add_field(p, "PAGE")


def add_cover(doc):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(90)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    run = kicker.add_run("TECHNICAL CHANGE GUIDE")
    set_run_font(run, size=11, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("Import Library")
    set_run_font(run, size=30, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(28)
    run = subtitle.add_run("Code changes, architecture, security and implementation analysis")
    set_run_font(run, size=14, color=DARK_BLUE)

    summary = doc.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary.paragraph_format.left_indent = Inches(0.55)
    summary.paragraph_format.right_indent = Inches(0.55)
    summary.paragraph_format.space_after = Pt(80)
    run = summary.add_run(
        "Chuyển chức năng import từ localStorage sang backend/Supabase, "
        "tạo bản sao độc lập của library, document, tag, AI chunks và file Storage."
    )
    set_run_font(run, size=11, color=MUTED, italic=True)

    meta = doc.add_table(rows=2, cols=2)
    meta.autofit = False
    widths = [Inches(3.25), Inches(3.25)]
    values = [
        ("Phạm vi", "4 file nguồn"),
        ("Định dạng", "Code + phân tích"),
        ("Kiểm tra", "Build, ESLint, syntax"),
        ("Trạng thái", "Đã triển khai"),
    ]
    for index, (label, value) in enumerate(values):
        row, col = divmod(index, 2)
        cell = meta.cell(row, col)
        cell.width = widths[col]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(5)
        label_run = p.add_run(f"{label}: ")
        set_run_font(label_run, size=9.5, color=DARK_BLUE, bold=True)
        value_run = p.add_run(value)
        set_run_font(value_run, size=9.5, color=NAVY)

    doc.add_page_break()


def is_code_start(line):
    stripped = line.strip()
    starters = (
        "async function ",
        "function ",
        "exports.",
        "const ",
        "let ",
        "if (",
        "return ",
        "router.",
        "import ",
        "export ",
        "<button",
        "<strong",
        "<i ",
        "{is",
        "set",
        "await ",
        "try {",
        "} catch",
        "} finally",
    )
    return stripped.startswith(starters)


def add_code_block(doc, lines):
    if not lines:
        return
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.14)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    set_paragraph_shading(p, CODE_FILL, CODE_BORDER)
    for index, line in enumerate(lines):
        run = p.add_run(line)
        set_run_font(run, name="Consolas", size=8.2, color="27313D")
        if index < len(lines) - 1:
            run.add_break()


def add_body_from_source(doc):
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    index = 2
    code_lines = []
    code_mode = False

    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()

        if code_mode:
            if not stripped:
                add_code_block(doc, code_lines)
                code_lines = []
                code_mode = False
                index += 1
                continue
            code_lines.append(line)
            index += 1
            continue

        if not stripped or set(stripped) <= {"=", "-"}:
            index += 1
            continue

        if is_code_start(line):
            code_mode = True
            code_lines.append(line)
            index += 1
            continue

        if re.match(r"^\d+\.\s+(BE/|FE/|LUỒNG|ĐÁNH GIÁ|KẾT QUẢ)", stripped):
            doc.add_paragraph(stripped, style="Heading 1")
        elif re.match(r"^[A-Z]\.\s+", stripped):
            doc.add_paragraph(stripped, style="Heading 2")
        elif (
            re.match(r"^\d+\.\s+", stripped)
            and len(stripped) < 90
            and not stripped.startswith(("1. Người", "2. Frontend", "3. Backend", "4. Backend", "5. Người", "6. Frontend", "7. Backend", "8. Frontend"))
        ):
            doc.add_paragraph(stripped, style="Heading 3")
        elif stripped.endswith(":") and len(stripped) < 80:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(stripped)
            set_run_font(run, size=11, color=DARK_BLUE, bold=True)
        elif stripped.startswith("- "):
            p = doc.add_paragraph(stripped[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(re.sub(r"^\d+\.\s+", "", stripped), style="List Number")
        else:
            p = doc.add_paragraph(stripped)
            p.paragraph_format.widow_control = True

        index += 1

    add_code_block(doc, code_lines)


def main():
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    add_cover(doc)
    add_body_from_source(doc)

    core = doc.core_properties
    core.title = "Import Library - Code Changes and Analysis"
    core.subject = "AI Study Hub technical implementation guide"
    core.author = "AI Study Hub"
    core.keywords = "Import Library, Supabase, React, Express, code analysis"

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
