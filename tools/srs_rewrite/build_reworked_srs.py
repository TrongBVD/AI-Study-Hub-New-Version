from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "output" / "srs"
OUT_PATH = OUT_DIR / "AI_StudyHub_SRS_v2.0_Reworked_IEEE.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 33, 61)
MUTED = RGBColor(89, 89, 89)
LIGHT_FILL = "F2F4F7"
BLUE_FILL = "E8EEF5"
WHITE = "FFFFFF"
BLACK_HEX = "000000"


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(row.cells[idx])
            row.cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_update_fields_on_open(doc):
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.insert(0, update_fields)
    update_fields.set(qn("w:val"), "true")


def add_toc_field(paragraph, levels="1-3"):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f' TOC \\o "{levels}" \\h \\z \\u '

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    display = OxmlElement("w:t")
    display.text = "Mục lục sẽ được cập nhật khi mở file trong Microsoft Word."

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(display)
    run._r.append(fld_end)


def add_table(doc, headers, rows, widths, header_fill=LIGHT_FILL, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    header = table.rows[0]
    set_repeat_table_header(header)
    for i, text in enumerate(headers):
        cell = header.cells[i]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(str(text))
        set_run_font(run, size=font_size, color=INK, bold=True)
    for row_data in rows:
        row = table.add_row()
        for i, text in enumerate(row_data):
            cell = row.cells[i]
            set_cell_shading(cell, WHITE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(str(text))
            set_run_font(run, size=font_size, color=RGBColor(0, 0, 0))
    doc.add_paragraph()
    return table


def add_para(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix) :])
        set_run_font(r2)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_run_font(run)


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    set_update_fields_on_open(doc)


def set_running_header_footer(section):
    header = section.header.paragraphs[0]
    header.text = ""
    left = header.add_run("AI StudyHub SRS v2.0")
    set_run_font(left, size=9, color=MUTED, bold=True)
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("AI StudyHub SRS v2.0")
    set_run_font(footer_run, size=9, color=MUTED)


def add_cover(doc):
    for _ in range(3):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SOFTWARE REQUIREMENTS SPECIFICATION")
    set_run_font(run, size=24, color=INK, bold=True)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("AI StudyHub - AI-Powered Learning and Knowledge Management System")
    set_run_font(run, size=14, color=MUTED, italic=True)
    meta = [
        ("Project code", "SWP391_G1"),
        ("Version", "2.0 - Reworked IEEE SRS"),
        ("Date", date.today().isoformat()),
        ("Prepared for", "SWP391 Capstone Project"),
        ("Prepared by", "Bui Vo Duc Trong, Ha Thuc Nhat Thanh, Vo Dang Khoa, Le Minh Hoang"),
    ]
    doc.add_paragraph()
    add_table(doc, ["Field", "Value"], meta, [1.6, 4.9], header_fill=BLUE_FILL, font_size=10)
    add_para(
        doc,
        "This document reorganizes the AI StudyHub requirements using the supplied contents structure and an IEEE-style SRS approach. It is grounded in the current React frontend, Express backend, Supabase data/storage layer, and Gemini AI integration observed in the project repository.",
    )
    doc.add_page_break()


def add_record_of_changes(doc):
    add_heading(doc, "I. Record of Changes", 1)
    rows = [
        ("Original", "A", "SWP391_G1", "Initial SRS with user stories, non-functional requirements, and database dictionary."),
        ("2026-06-12", "M", "SWP391_G1", "Added System Admin clarification, admin requirements, workflows, API endpoints, and database updates."),
        (date.today().isoformat(), "M", "Codex / BA Review", "Reworked SRS to follow the supplied contents structure and IEEE-style requirement organization; aligned requirements with current repo evidence."),
    ]
    add_table(doc, ["Date", "A/M/D", "In charge", "Change Description"], rows, [1.1, 0.7, 1.8, 2.9], font_size=9.2)
    add_para(doc, "*A - Added, M - Modified, D - Deleted.")
    doc.add_page_break()


def add_contents(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(18)
    run = title.add_run("MỤC LỤC")
    set_run_font(run, size=18, color=INK, bold=True)

    toc = doc.add_paragraph()
    toc.paragraph_format.space_after = Pt(0)
    add_toc_field(toc, levels="1-3")
    doc.add_page_break()


def add_product_overview(doc):
    add_heading(doc, "II. Software Requirement Specification", 1)
    add_heading(doc, "1. Product Overview", 1)
    add_heading(doc, "1.1 Product Purpose", 2)
    add_para(
        doc,
        "AI StudyHub is a web-based learning knowledge platform for students who need one place to store academic documents, organize study libraries and workspaces, search learning materials, and ask AI-assisted questions over uploaded content. The system reduces document fragmentation across chat groups, drives, email, and local storage by centralizing study materials and adding controlled AI processing.",
    )
    add_heading(doc, "1.2 Product Scope", 2)
    add_bullets(
        doc,
        [
            "Account access through Google SSO, profile completion, email/OTP verification, username/password login, logout, and password reset.",
            "Personal library and document management, including file upload, metadata storage, download, soft delete, and document listing.",
            "AI moderation of uploaded study documents, text extraction, chunking, embeddings, RAG chatbot answers, and AI-generated flashcards.",
            "Workspace collaboration, including workspace creation, member listing, user search, and assignment of workspace-level roles.",
            "System Admin module for dashboard metrics, AI moderation review, user status management, activity logs, and usage monitoring.",
            "Quota and abuse controls for upload/download volume and AI chatbot usage.",
        ],
    )
    add_heading(doc, "1.3 Product Perspective", 2)
    add_para(
        doc,
        "The product is implemented as a React/Vite single-page web application that communicates with a Node.js/Express API. Supabase is used for relational data, pgvector-backed retrieval, and document storage. Google Gemini provides text generation, document moderation, embeddings, contextual question answering, and flashcard generation.",
    )
    add_heading(doc, "1.4 Users and Stakeholders", 2)
    rows = [
        ("Student / User", "Primary user who uploads, organizes, searches, downloads, and studies documents with AI assistance."),
        ("Workspace Admin", "Workspace-scoped user who creates a workspace and manages members within that workspace only."),
        ("System Admin", "Platform administrator with role SYSTEM_ADMIN who accesses /admin and /api/admin functions."),
        ("Development Team", "Implements frontend, backend, data model, AI integration, security, and deployment."),
        ("Instructor / Evaluator", "Reviews requirement clarity, scope consistency, and evidence of implementation alignment."),
    ]
    add_table(doc, ["Stakeholder", "Interest"], rows, [1.8, 4.7], font_size=9.5)
    add_heading(doc, "1.5 Operating Environment", 2)
    rows = [
        ("Client", "Modern web browser with JavaScript support; React frontend served by Vite in development."),
        ("Server", "Node.js/Express backend exposing REST APIs under /api/auth, /api/documents, /api/ai, /api/admin, and /api/workspaces."),
        ("Database", "Supabase PostgreSQL with tables for profiles, otp_tokens, documents, tags, document_chunks, flashcards, workspaces, workspace_members, activity_logs, daily_quota_usage, ai_usage_logs, and libraries."),
        ("Storage", "Supabase Storage bucket for uploaded document binaries; signed URLs are used for download access."),
        ("AI Services", "Google Gemini text and embedding models; pgvector RPC match_document_chunks for semantic retrieval."),
    ]
    add_table(doc, ["Layer", "Environment"], rows, [1.4, 5.1], font_size=9.1)
    add_heading(doc, "1.6 Constraints, Assumptions, and Dependencies", 2)
    add_bullets(
        doc,
        [
            "Only authenticated users may access protected dashboard, document, AI, and workspace functions.",
            "Only users with role SYSTEM_ADMIN may access Admin pages and Admin APIs.",
            "Workspace Admin is not the same role as System Admin; Workspace Admin is scoped to a single workspace.",
            "AI outcomes are probabilistic and must fail safely; document processing errors must not expose stack traces to end users.",
            "Mobile application requirements are not implemented in the current repository and are treated as future scope unless the team adds a mobile codebase.",
        ],
    )


def add_user_requirements(doc):
    add_heading(doc, "2. User Requirements", 1)
    add_heading(doc, "2.1 Actors", 2)
    add_para(
        doc,
        "Primary actors are human users who directly initiate system goals. Secondary actors are automated services or external platforms that support those goals.",
    )
    rows = [
        ("A01", "Guest", "Primary", "Unauthenticated visitor who can open landing, login, registration, OTP, and password recovery screens."),
        ("A02", "Student / Registered User", "Primary", "Authenticated learner who manages profile, libraries, documents, workspaces, search, flashcards, and AI chat."),
        ("A03", "Workspace Admin", "Primary", "Workspace-scoped member role. Can create a workspace and manage workspace members; this is not platform administration."),
        ("A04", "Workspace Editor", "Primary", "Workspace member intended to contribute or manage content according to workspace permission policy."),
        ("A05", "Workspace Viewer", "Primary", "Workspace member intended to view shared workspace content with restricted modification rights."),
        ("A06", "System Admin", "Primary", "Platform-level administrator with SYSTEM_ADMIN role for moderation, users, logs, usage, and dashboard metrics."),
        ("A07", "AI Engine", "Secondary", "Automated Gemini-backed service that moderates documents, creates embeddings, answers with context, and generates flashcards."),
        ("A08", "External Identity/Email Provider", "Secondary", "Google OAuth and email/OTP delivery services used during authentication and password recovery."),
        ("A09", "Supabase Platform", "Secondary", "External database and storage platform used for metadata, storage, signed URLs, and vector search RPC."),
    ]
    add_table(doc, ["ID", "Actor", "Type", "Description"], rows, [0.5, 1.45, 0.85, 3.7], font_size=8.5)
    add_heading(doc, "2.2 Use Cases", 2)
    add_para(
        doc,
        "The primary use-case view is web application based. Workspace member roles specialize the registered user actor, while System Admin remains a separate platform-level actor.",
    )
    use_cases = [
        ("UC-01", "Register with Google and complete profile", "Guest", "Create a profile through Google SSO, OTP verification, and username/password completion."),
        ("UC-02", "Log in and log out", "Student, System Admin", "Authenticate with credentials and receive a JWT/session; logout clears client access."),
        ("UC-03", "Recover password with OTP", "Guest", "Request OTP, verify it before expiry, and set a new password."),
        ("UC-04", "Manage personal profile", "Student", "View profile data and update supported public/private profile fields."),
        ("UC-05", "Create and manage libraries", "Student", "Create a personal library, edit metadata, and navigate library content."),
        ("UC-06", "Upload study documents", "Student", "Upload PDF/Word files subject to type, size, quota, and ownership rules."),
        ("UC-07", "AI-moderate uploaded content", "AI Engine", "Extract text, classify document suitability, store status and AI rejection reason."),
        ("UC-08", "Download allowed document", "Student", "Retrieve a signed download URL when ownership/public/workspace permissions allow access."),
        ("UC-09", "Delete a document", "Student", "Soft-delete a document owned by the user."),
        ("UC-10", "Ask a document question", "Student, AI Engine", "Submit a question, retrieve relevant chunks, and return a contextual answer with source chunk hints."),
        ("UC-11", "Generate flashcards", "Student, AI Engine", "Generate 5-10 study cards from approved document chunks and save them."),
        ("UC-12", "Search users", "Student, Workspace Admin", "Find existing users by username/profile information for profile viewing or workspace membership."),
        ("UC-13", "Create workspace", "Student", "Create a workspace and automatically become workspace Admin."),
        ("UC-14", "Manage workspace members", "Workspace Admin", "Search users, add members, and assign Editor or Viewer roles."),
        ("UC-15", "Access workspace details", "Workspace member", "View workspace metadata and membership if the requester belongs to the workspace."),
        ("UC-16", "View admin dashboard", "System Admin", "View platform counts, moderation queue volume, quota usage, and AI usage."),
        ("UC-17", "Review AI-blocked documents", "System Admin", "Approve or keep rejected AI-flagged documents with an audit reason."),
        ("UC-18", "Manage user status", "System Admin", "Disable or reactivate user accounts, excluding unsafe self-demotion/self-disable behavior."),
        ("UC-19", "View activity logs", "System Admin", "Filter and inspect immutable activity logs for moderation, user, document, and quota actions."),
        ("UC-20", "Monitor usage", "System Admin", "Review daily upload, download, chat count, and token usage by user/date."),
        ("UC-21", "Use mobile application", "Future mobile actor", "Not in current repository; reserved for future mobile-specific requirements if a mobile app is added."),
    ]
    add_table(doc, ["ID", "Use Case", "Actors", "Use Case Description"], use_cases, [0.65, 1.85, 1.45, 2.55], font_size=8.1)


def add_functional_requirements(doc):
    add_heading(doc, "3. Functional Requirements", 1)
    add_heading(doc, "3.1 System Functional Overview", 2)
    add_heading(doc, "3.1.1 Web Screen Flow", 3)
    add_para(
        doc,
        "The web flow starts at the landing page and authentication screens. After login, regular users enter /dashboard and can access home, libraries, library detail, import/upload, settings, create workspace, workspace list, workspace detail, profile, flashcards, user search, and document search. System Admin users enter /admin and can access dashboard, moderation, user management, logs, usage, settings, and profile.",
    )
    screen_rows = [
        ("Landing", "/", "Guest entry page and navigation to authentication."),
        ("Login", "/login", "Credential-based login for students and System Admins."),
        ("Register / OTP / Complete Profile", "/register, /verify-otp, /complete-profile", "Google registration, OTP verification, and profile setup."),
        ("Password Recovery", "/forgot-password, /reset-password", "OTP-based password reset flow."),
        ("Dashboard Home", "/dashboard/home", "Authenticated user dashboard landing area."),
        ("Libraries", "/dashboard/libraries", "List personal libraries and documents."),
        ("Import Library / Upload", "/dashboard/import-library", "Upload document files and pass workspaceId when applicable."),
        ("Library Detail", "/dashboard/libraries/:libraryId", "View a library and its document content."),
        ("Workspace List / Create", "/dashboard/workspaces, /dashboard/create-workspace", "List and create collaborative workspaces."),
        ("Workspace Detail", "/dashboard/workspaces/:workspaceId", "View workspace details, members, search users, and add members."),
        ("Profile", "/dashboard/profile, /dashboard/profile/:id", "View own or other public profile details."),
        ("Flashcards", "/dashboard/flashcards", "Study AI-generated flashcards."),
        ("Search", "/dashboard/search", "Search for documents or learning materials."),
        ("Admin Dashboard", "/admin/dashboard", "System metrics, moderation counts, quota and AI usage totals."),
        ("Admin Moderation", "/admin/moderation", "AI-blocked document review."),
        ("Admin Users", "/admin/users", "User search and account status management."),
        ("Admin Logs", "/admin/logs", "Activity log review and filtering."),
        ("Admin Usage", "/admin/usage", "Upload/download/chat/token usage monitoring."),
    ]
    add_table(doc, ["Screen", "Route", "Description"], screen_rows, [1.45, 2.05, 3.0], font_size=8.3)
    add_heading(doc, "3.1.2 Backend Service Overview", 3)
    endpoint_rows = [
        ("Auth", "/api/auth", "Google login, OTP, complete setup, login/logout, forgot/reset password, user search, public profile lookup."),
        ("Documents", "/api/documents", "List, upload, download, soft delete documents; create/update libraries."),
        ("AI", "/api/ai", "Chat with document and generate flashcards for allowed approved documents."),
        ("Workspaces", "/api/workspaces", "List/create workspace, get detail, list members, search users, add members."),
        ("Admin", "/api/admin", "Dashboard, moderation review, users, logs, and usage monitoring."),
    ]
    add_table(doc, ["Domain", "API Root", "Responsibility"], endpoint_rows, [1.2, 1.45, 3.85], font_size=9.0)
    add_heading(doc, "3.1.3 Requirement Priority Definitions", 3)
    add_table(
        doc,
        ["Priority", "Meaning"],
        [
            ("Must", "Required for MVP correctness, security, or primary user workflow."),
            ("Should", "Important for quality, usability, or operational control but not a hard launch blocker."),
            ("Could", "Useful enhancement or future iteration."),
        ],
        [1.0, 5.5],
        font_size=9.5,
    )

    add_heading(doc, "3.2 Web Application", 2)
    reqs = [
        ("FR-AUTH-01", "Auth", "Must", "The system shall allow users to register or log in using Google OAuth.", "A valid Google token creates or locates a profile and returns public user data/session information."),
        ("FR-AUTH-02", "Auth", "Must", "The system shall support OTP verification during profile completion and password recovery.", "OTP is matched by email/code, checked against expiry, and deleted after use."),
        ("FR-AUTH-03", "Auth", "Must", "The system shall allow username/password login and logout.", "Valid credentials return an access token; logout clears the client-side session."),
        ("FR-AUTH-04", "Auth", "Must", "The system shall store passwords as bcrypt hashes.", "No plaintext password is returned by API responses or displayed to Admin users."),
        ("FR-AUTH-05", "Auth", "Must", "Protected pages and APIs shall require authenticated access.", "Unauthenticated access to /dashboard and protected API routes is redirected or rejected."),
        ("FR-DOC-01", "Document", "Must", "The system shall allow authenticated users to upload supported study document files.", "Upload rejects missing files, unsupported types, and daily upload quota breaches."),
        ("FR-DOC-02", "Document", "Must", "The system shall store uploaded files in Supabase Storage and metadata in documents.", "Each uploaded file receives uploader_id, title, file_url, file_size_bytes, is_public, workspace_id when supplied, and status."),
        ("FR-DOC-03", "Document", "Must", "The system shall extract document text, split it into chunks, and store embeddings.", "Readable files create document_chunks rows with chunk_index, content, and 768-dimension vector literals."),
        ("FR-DOC-04", "Document", "Must", "The system shall moderate uploaded documents with AI before approval.", "Invalid or unreadable files are rejected with ai_reject_reason; approved files become available by permission rules."),
        ("FR-DOC-05", "Document", "Must", "The system shall allow document download only when the requester has permission and quota.", "Owner/public allowed documents return a short-lived signed URL; unauthorized access is rejected."),
        ("FR-DOC-06", "Document", "Must", "The system shall support soft deletion of user documents.", "Deleting a document sets deleted_at and removes it from normal listing/search views."),
        ("FR-LIB-01", "Library", "Should", "The system shall allow users to create and update personal libraries.", "Library metadata can be created and edited for the authenticated user."),
        ("FR-AI-01", "AI Chat", "Must", "The system shall answer user questions using only approved allowed document content.", "Chat rejects non-approved, missing, or forbidden documents and returns answer plus source chunk hints."),
        ("FR-AI-02", "AI Chat", "Must", "The system shall enforce a daily AI chatbot quota.", "After 50 chats per user/day, AI chat returns a quota exceeded response."),
        ("FR-AI-03", "Flashcards", "Must", "The system shall generate flashcards from approved document chunks.", "AI returns 5-10 question/answer cards that are saved to flashcards for the document."),
        ("FR-AI-04", "AI Processing", "Should", "The system shall generate or normalize tags when AI tag generation is available.", "Generated tags are inserted in tags and associated through document_tags without breaking upload if tag generation fails."),
        ("FR-SEARCH-01", "Search", "Should", "The system shall support document/user search interfaces.", "Users can search documents from dashboard search and search existing users for profile/workspace workflows."),
        ("FR-WS-01", "Workspace", "Must", "The system shall allow authenticated users to create workspaces.", "The creator is inserted into workspace_members with role Admin."),
        ("FR-WS-02", "Workspace", "Must", "The system shall list only workspaces where the requester is a member.", "Deleted workspaces are excluded and the requester role is returned as myRole."),
        ("FR-WS-03", "Workspace", "Must", "The system shall allow workspace members to view workspace detail and member list.", "Non-members cannot access workspace detail or member data."),
        ("FR-WS-04", "Workspace", "Must", "Workspace Admin shall be able to add existing users as Editor or Viewer.", "Adding a nonexistent user, duplicate member, or unsupported role is rejected."),
        ("FR-ADM-01", "Admin", "Must", "Only SYSTEM_ADMIN shall access Admin pages and APIs.", "Frontend route guard and backend requireAdmin middleware enforce System Admin role."),
        ("FR-ADM-02", "Admin", "Must", "The system shall provide dashboard metrics.", "Admin dashboard returns total users, total documents, pending moderation count, daily upload/download totals, AI chats, and tokens."),
        ("FR-ADM-03", "Admin", "Must", "System Admin shall review rejected, flagged, and retry-pending documents.", "Admin can approve or keep rejected with reason; action updates document status and audit data."),
        ("FR-ADM-04", "Admin", "Must", "System Admin shall manage user account status.", "Admin can disable/reactivate accounts and the action is logged."),
        ("FR-ADM-05", "Admin", "Should", "System Admin shall view immutable activity logs.", "Logs can be filtered and include actor, action, entity, old data, new data, timestamp, IP, and user agent when available."),
        ("FR-ADM-06", "Admin", "Should", "System Admin shall monitor daily quota and AI usage.", "Usage view combines daily_quota_usage and ai_usage_logs with user metadata."),
        ("FR-AUD-01", "Audit", "Must", "Important admin and moderation actions shall be written to activity_logs.", "Status changes, moderation decisions, user status updates, and quota interventions produce audit records."),
        ("FR-NOT-01", "Notification", "Could", "The system should notify users about moderation decisions and quota warnings.", "Notifications are represented in the data model but should be verified against implementation before claiming full delivery."),
    ]
    add_table(doc, ["ID", "Module", "Priority", "Requirement", "Acceptance Criteria"], reqs, [0.85, 0.9, 0.65, 2.2, 1.9], font_size=7.5)

    add_heading(doc, "3.3 Mobile Application", 2)
    add_para(
        doc,
        "No mobile application codebase is present in the current project repository. To keep the SRS accurate, mobile application requirements are classified as future scope. If a mobile app is later added, this section should be expanded with mobile actors, screen flow, offline behavior, device permissions, push notifications, and platform-specific acceptance criteria.",
    )
    add_table(
        doc,
        ["Mobile Area", "Current Requirement Status"],
        [
            ("Mobile authentication", "Not applicable to current repo; reuse web authentication requirements if mobile is introduced."),
            ("Mobile document upload/download", "Future scope; must define file picker, storage permissions, and quota behavior."),
            ("Mobile AI chat/flashcards", "Future scope; must define mobile UX, loading states, and rate-limit display."),
            ("Mobile workspace collaboration", "Future scope; must define member management and role restrictions on small screens."),
        ],
        [2.0, 4.5],
        font_size=9.0,
    )


def add_non_functional(doc):
    add_heading(doc, "4. Non-Functional Requirements", 1)
    add_heading(doc, "4.1 External Interfaces", 2)
    rows = [
        ("User Interface", "React web UI with protected routes, dashboard layout, admin layout, forms, tables, search, and confirmation dialogs."),
        ("REST API", "JSON APIs under /api/auth, /api/documents, /api/ai, /api/admin, and /api/workspaces."),
        ("Database", "Supabase PostgreSQL accessed through @supabase/supabase-js; pgvector RPC for semantic chunk matching."),
        ("Storage", "Supabase Storage bucket for uploaded files and signed URL generation for downloads."),
        ("AI Provider", "Google Gemini APIs through @google/genai for text generation and embeddings."),
        ("Email/OAuth", "Google OAuth and nodemailer-backed OTP/password recovery flows."),
    ]
    add_table(doc, ["Interface", "Requirement"], rows, [1.4, 5.1], font_size=9.0)
    add_heading(doc, "4.2 Quality Attributes", 2)
    rows = [
        ("Performance", "Search responses should complete within 2 seconds for normal MVP data volume; Admin dashboard should load within 3 seconds; AI chat should respond within 5 seconds under normal provider latency."),
        ("Security", "All protected endpoints require authentication; admin endpoints require SYSTEM_ADMIN; passwords are bcrypt hashed; JWT/session expiry must be enforced."),
        ("Privacy", "Private documents and workspace documents are visible only to authorized users; Admins do not see password hashes or hidden private profile fields in UI."),
        ("Reliability", "AI failures, invalid JSON, missing chunks, and Supabase errors must return safe errors without crashing the server."),
        ("Availability", "The system targets cloud-hosted availability consistent with Supabase/provider uptime; user-facing errors must be clear during service degradation."),
        ("Auditability", "Admin write actions and moderation decisions must create activity log records with enough metadata to investigate behavior."),
        ("Usability", "Dangerous actions require confirmation; error/success messages use plain language; tables support filtering or pagination where data can grow."),
        ("Maintainability", "Backend routing, controllers, services, middleware, and Supabase configuration remain separated; frontend API utilities isolate API calls from page components."),
        ("Scalability", "Document chunks and embeddings support retrieval growth; daily quota tables and usage logs support cost control."),
        ("Compatibility", "The web app supports modern Chromium/Firefox browsers with JavaScript enabled."),
    ]
    add_table(doc, ["Attribute", "Requirement"], rows, [1.25, 5.25], font_size=8.8)


def add_appendix(doc):
    add_heading(doc, "5. Requirement Appendix", 1)
    add_heading(doc, "5.1 Business Rules", 2)
    rows = [
        ("BR-01", "Only Gmail/valid email accounts can complete supported registration and OTP flows."),
        ("BR-02", "OTP codes must expire and must not be reused after successful verification."),
        ("BR-03", "Password reset must require a valid OTP before password update."),
        ("BR-04", "Every uploaded document initially enters a pending/processing state before AI moderation result is applied."),
        ("BR-05", "Documents with insufficient readable text are rejected with a clear reason."),
        ("BR-06", "Documents rejected or flagged by AI can be reviewed by System Admin."),
        ("BR-07", "Only document owners or authorized users can download non-public documents."),
        ("BR-08", "AI chat and flashcard generation require an approved document and document chunks."),
        ("BR-09", "Daily AI chat count is limited to 50 questions per user unless changed by configuration."),
        ("BR-10", "Daily upload and download usage must be tracked for quota monitoring."),
        ("BR-11", "Workspace creator becomes workspace Admin for that workspace."),
        ("BR-12", "Workspace Admin permissions do not grant System Admin access."),
        ("BR-13", "System Admin access requires profile role SYSTEM_ADMIN."),
        ("BR-14", "System Admin decisions must record a reason when reviewing moderation outcomes."),
        ("BR-15", "Activity logs are append-only from the user interface perspective."),
        ("BR-16", "Soft-deleted documents are hidden from normal listings and downloads."),
        ("BR-17", "Admin users must not disable their own account through normal user status management."),
        ("BR-18", "AI tag generation failure must not block the core upload flow when moderation and chunking succeed."),
    ]
    add_table(doc, ["ID", "Business Rule"], rows, [0.8, 5.7], font_size=8.8)

    add_heading(doc, "5.2 Common Requirements", 2)
    rows = [
        ("Authentication", "Protected APIs validate bearer token/session and reject missing, expired, or invalid tokens."),
        ("Authorization", "Role checks are enforced in backend middleware, not only by frontend route guards."),
        ("Validation", "Required IDs, emails, file presence, file type, role values, and decisions are validated before data changes."),
        ("Error Handling", "Errors return structured JSON with status and message; internal errors are logged server-side."),
        ("Pagination/Filtering", "Admin user lists, logs, usage, and moderation queues should support filtering and pagination as data grows."),
        ("Audit", "Write actions that affect users, documents, moderation, quota, or workspace membership create activity logs."),
        ("Data Deletion", "Soft delete is preferred for user-facing document deletion; permanent deletion should be scheduled and auditable."),
        ("AI Safety", "AI results are parsed defensively; invalid JSON or provider timeout returns a controlled error."),
        ("Traceability", "Each functional requirement should map to at least one page, API endpoint, service, or table before implementation sign-off."),
    ]
    add_table(doc, ["Common Area", "Requirement"], rows, [1.55, 4.95], font_size=9.0)

    add_heading(doc, "5.3 Application Messages List", 2)
    rows = [
        ("AUTH-001", "Token Google không hợp lệ / Invalid Google token.", "Google login fails verification."),
        ("AUTH-002", "Mã OTP không hợp lệ hoặc đã hết hạn.", "OTP not found, expired, or already consumed."),
        ("AUTH-003", "Tài khoản không ở trạng thái chờ hoàn tất hồ sơ.", "OTP verification is attempted for an account not pending setup."),
        ("AUTH-004", "Invalid username/email or password.", "Credential login fails."),
        ("AUTH-005", "Your account is disabled.", "Disabled profile attempts protected access."),
        ("DOC-001", "No files were uploaded.", "Upload request has no files."),
        ("DOC-002", "Daily upload quota exceeded.", "Upload would exceed configured daily bytes."),
        ("DOC-003", "Could not extract enough readable text from this file.", "Text extraction produces insufficient content."),
        ("DOC-004", "You do not have permission to download this document.", "Download permission check fails."),
        ("DOC-005", "Document not found.", "Document ID is missing, deleted, or inaccessible."),
        ("AI-001", "This document is not approved or not ready for AI chat yet.", "AI chat requested before document approval."),
        ("AI-002", "No AI chunks found for this document. Re-upload or re-process it.", "Vector chunks are missing."),
        ("AI-003", "Daily AI chatbot quota exceeded.", "User has reached the daily chat limit."),
        ("AI-004", "AI did not return valid JSON.", "Gemini response cannot be parsed."),
        ("WS-001", "Workspace name is required.", "Workspace creation missing name."),
        ("WS-002", "You cannot access this workspace.", "Requester is not a workspace member."),
        ("WS-003", "Role must be Editor or Viewer.", "Workspace Admin attempts unsupported member role."),
        ("ADM-001", "Admin access required.", "Non-admin attempts admin page/API access."),
        ("ADM-002", "decision must be APPROVE or KEEP_REJECTED.", "Invalid moderation decision submitted."),
        ("ADM-003", "Reason is required.", "Admin moderation decision missing reason."),
        ("ADM-004", "Could not load admin dashboard.", "Dashboard metrics query fails."),
        ("GEN-001", "Something went wrong. Please try again.", "Generic fallback message for unexpected UI/API failures."),
    ]
    add_table(doc, ["Code", "Message", "Trigger"], rows, [0.85, 2.45, 3.2], font_size=8.1)

    add_heading(doc, "5.4 Other Requirements", 2)
    add_heading(doc, "5.4.1 Data Model Summary", 3)
    rows = [
        ("profiles", "User identity, email, username, full_name, role, status, profile privacy, login/session metadata."),
        ("otp_tokens", "Email OTP code, expiry, and one-time verification records."),
        ("documents", "Document metadata, owner, workspace, storage path, size, public flag, status, AI/admin review fields, deletion time."),
        ("tags / document_tags", "Normalized tags and many-to-many association with documents."),
        ("document_chunks", "Chunked document text with 768-dimension pgvector embeddings for RAG/search."),
        ("flashcards", "AI-generated study question/answer cards linked to documents."),
        ("workspaces / workspace_members", "Workspace metadata and membership roles Admin, Editor, Viewer."),
        ("activity_logs", "Audit records for important user/admin/system actions."),
        ("daily_quota_usage", "Per-user per-day upload and download byte counters."),
        ("ai_usage_logs", "Per-user per-day token and chatbot usage counters."),
        ("libraries", "Personal library metadata used by the frontend library pages."),
    ]
    add_table(doc, ["Table", "Purpose"], rows, [1.55, 4.95], font_size=8.6)
    add_heading(doc, "5.4.2 API Endpoint Summary", 3)
    rows = [
        ("POST", "/api/auth/google", "Google login."),
        ("POST", "/api/auth/verify-otp", "Verify profile completion OTP."),
        ("GET", "/api/auth/check-username", "Check username availability."),
        ("POST", "/api/auth/complete-setup", "Complete user profile setup."),
        ("POST", "/api/auth/login", "Credential login."),
        ("POST", "/api/auth/logout", "Logout."),
        ("POST", "/api/auth/forgot-password", "Request password reset OTP."),
        ("POST", "/api/auth/verify-reset-otp", "Verify reset OTP."),
        ("POST", "/api/auth/reset-password", "Set new password."),
        ("GET", "/api/auth/search", "Search users."),
        ("GET", "/api/auth/users/:id/profile", "View public profile."),
        ("GET", "/api/documents", "List current user's documents."),
        ("POST", "/api/documents/upload", "Upload up to 10 files."),
        ("GET", "/api/documents/:documentId/download", "Create signed download URL."),
        ("DELETE", "/api/documents/:documentId", "Soft delete document."),
        ("POST", "/api/ai/chat", "Ask a document question."),
        ("POST", "/api/ai/documents/:documentId/flashcards", "Generate flashcards."),
        ("GET", "/api/workspaces", "List my workspaces."),
        ("POST", "/api/workspaces", "Create workspace."),
        ("GET", "/api/workspaces/:workspaceId", "Get workspace detail."),
        ("GET", "/api/workspaces/:workspaceId/members", "List workspace members."),
        ("GET", "/api/workspaces/:workspaceId/users/search", "Search users for workspace."),
        ("POST", "/api/workspaces/:workspaceId/members", "Add workspace member."),
        ("GET", "/api/admin/dashboard", "Admin dashboard metrics."),
        ("GET", "/api/admin/moderation", "Moderation document queue."),
        ("PATCH", "/api/admin/moderation/:documentId", "Review moderated document."),
        ("GET", "/api/admin/users", "List/search users."),
        ("PATCH", "/api/admin/users/:userId/status", "Update user status."),
        ("GET", "/api/admin/logs", "View activity logs."),
        ("GET", "/api/admin/usage", "View quota and AI usage."),
    ]
    add_table(doc, ["Method", "Endpoint", "Purpose"], rows, [0.75, 2.55, 3.2], font_size=7.7)
    add_heading(doc, "5.4.3 Traceability Matrix", 3)
    rows = [
        ("Authentication", "UC-01, UC-02, UC-03", "FR-AUTH-01 through FR-AUTH-05", "App.jsx auth routes, authRoutes.js, authController.js, authService.js"),
        ("Document Management", "UC-05, UC-06, UC-08, UC-09", "FR-DOC-01 through FR-DOC-06, FR-LIB-01", "documentRoutes.js, documentController.js, documentApi.js"),
        ("AI Study Assistance", "UC-07, UC-10, UC-11", "FR-AI-01 through FR-AI-04", "aiRoutes.js, aiController.js, aiService.js"),
        ("Workspace Collaboration", "UC-13, UC-14, UC-15", "FR-WS-01 through FR-WS-04", "workspaceRoutes.js, workspaceController.js, workspaceApi.js"),
        ("System Administration", "UC-16 through UC-20", "FR-ADM-01 through FR-ADM-06, FR-AUD-01", "adminRoutes.js, adminController.js, requireAdmin.js, Admin pages"),
        ("Mobile", "UC-21", "Future scope", "No current mobile codebase."),
    ]
    add_table(doc, ["Area", "Use Cases", "Requirements", "Evidence"], rows, [1.1, 1.35, 1.65, 2.4], font_size=7.8)
    add_heading(doc, "5.4.4 Open BA Questions", 3)
    add_bullets(
        doc,
        [
            "Should the final SRS remain fully English, or should the team submit a bilingual English/Vietnamese version?",
            "Should mobile application requirements be removed entirely from the required contents, or retained as future scope as written here?",
            "What exact production quotas should be configured for upload/download if the 50 MB/day values change before submission?",
            "Should Workspace Editor be allowed to upload/delete shared workspace documents, or only view/share according to a narrower permission model?",
            "Should notifications be implemented for moderation outcomes and workspace invitations before final SRS sign-off?",
        ],
    )


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    set_running_header_footer(doc.sections[0])
    add_cover(doc)
    add_record_of_changes(doc)
    add_contents(doc)
    add_product_overview(doc)
    add_user_requirements(doc)
    add_functional_requirements(doc)
    add_non_functional(doc)
    add_appendix(doc)

    for section in doc.sections:
        set_running_header_footer(section)

    doc.core_properties.title = "AI StudyHub Software Requirements Specification"
    doc.core_properties.subject = "IEEE-style SRS for AI StudyHub"
    doc.core_properties.author = "SWP391_G1 / Codex BA Rewrite"
    doc.core_properties.comments = "Reworked from source SRS, reference PDF contents template, and current project repository evidence."
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
