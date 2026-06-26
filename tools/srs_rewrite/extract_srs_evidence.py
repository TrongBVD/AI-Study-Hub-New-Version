from __future__ import annotations

import json
import re
from pathlib import Path

import docx
import pdfplumber


ROOT = Path(__file__).resolve().parents[2]
DOCX_PATH = ROOT / "AI_StudyHub_SRS_v2.0.source.docx"
PDF_PATH = ROOT / "reference_srs_report3.source.pdf"


def para_text(paragraph) -> str:
    return " ".join(paragraph.text.split())


def extract_docx() -> dict:
    document = docx.Document(DOCX_PATH)
    paragraphs = []
    for i, para in enumerate(document.paragraphs, start=1):
        text = para_text(para)
        if text:
            paragraphs.append(
                {
                    "index": i,
                    "style": para.style.name if para.style else "",
                    "text": text,
                }
            )

    tables = []
    for table_index, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            cells = [" ".join(cell.text.split()) for cell in row.cells]
            if any(cells):
                rows.append(cells)
        tables.append({"index": table_index, "rows": rows[:40]})

    return {
        "paragraph_count": len(paragraphs),
        "table_count": len(tables),
        "first_paragraphs": paragraphs[:140],
        "heading_like": [
            p
            for p in paragraphs
            if p["style"].lower().startswith("heading")
            or re.match(r"^\d+(\.\d+)*\s+\S", p["text"])
        ][:220],
        "tables": tables[:20],
    }


def extract_pdf() -> dict:
    pages = []
    with pdfplumber.open(PDF_PATH) as pdf:
        for page_number, page in enumerate(pdf.pages[:12], start=1):
            text = page.extract_text() or ""
            lines = [" ".join(line.split()) for line in text.splitlines()]
            lines = [line for line in lines if line]
            pages.append({"page": page_number, "lines": lines[:80]})
    return {"pages_read": len(pages), "pages": pages}


def scan_repo() -> dict:
    files = {}
    interesting = [
        "README.md",
        "BE/server.js",
        "BE/package.json",
        "FE/package.json",
        "FE/src/App.jsx",
        "BE/src/routes/authRoutes.js",
        "BE/src/routes/documentRoutes.js",
        "BE/src/routes/workspaceRoutes.js",
        "BE/src/routes/aiRoutes.js",
        "BE/src/routes/adminRoutes.js",
        "BE/src/controllers/authController.js",
        "BE/src/controllers/documentController.js",
        "BE/src/controllers/workspaceController.js",
        "BE/src/controllers/aiController.js",
        "BE/src/controllers/adminController.js",
        "BE/src/services/aiService.js",
        "BE/src/services/textExtractService.js",
        "BE/src/middleware/authMiddleware.js",
        "BE/src/middleware/requireAdmin.js",
        "FE/src/utils/api.js",
        "FE/src/utils/documentApi.js",
        "FE/src/utils/workspaceApi.js",
        "FE/src/utils/aiApi.js",
        "FE/src/utils/adminApi.js",
    ]
    for rel in interesting:
        path = ROOT / rel
        if path.exists():
            files[rel] = path.read_text(encoding="utf-8", errors="replace")

    route_pattern = re.compile(r"router\.(get|post|put|patch|delete)\(([^;]+)", re.I)
    routes = {}
    for rel, text in files.items():
        if "/routes/" in rel.replace("\\", "/"):
            routes[rel] = [m.group(0)[:220] for m in route_pattern.finditer(text)]

    frontend_pages = sorted(
        str(path.relative_to(ROOT)).replace("\\", "/")
        for path in (ROOT / "FE/src/components/pages").rglob("*.jsx")
    )

    return {"routes": routes, "frontend_pages": frontend_pages, "files": files}


def main() -> None:
    evidence = {
        "docx": extract_docx(),
        "pdf": extract_pdf(),
        "repo": scan_repo(),
    }
    out_path = ROOT / "tools/srs_rewrite/evidence.json"
    out_path.write_text(json.dumps(evidence, indent=2, ensure_ascii=False), encoding="utf-8")
    print(out_path)


if __name__ == "__main__":
    main()
