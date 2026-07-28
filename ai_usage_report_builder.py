from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT_FILE = r"D:\School materials\SWP391\AI-student-hub\AI_Usage_Report.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"].font.size = Pt(11)

for name, size in [("Title", 26), ("Heading 1", 16), ("Heading 2", 13)]:
    style = styles[name]
    style.font.name = "Arial"
    style.font.color.rgb = RGBColor(0, 0, 0)
    style.font.size = Pt(size)
    if name == "Title":
        style.font.bold = False

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = title.add_run("AI Usage Report for the AI Student Hub Project")
run.font.name = "Arial"
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0, 0, 0)
run.font.bold = False
title.paragraph_format.space_after = Pt(3)
title.paragraph_format.space_before = Pt(0)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(14)
subtitle.add_run("Prepared from project commit history and change patterns").italic = True

intro = doc.add_paragraph(
    "This report summarizes where AI-assisted work appears to have been used in the project, "
    "based on the change history and the kinds of fixes made. It focuses on practical development "
    "usage rather than model telemetry, which is not available in the repository."
)

sections = [
    ("Project Overview", [
        "AI was used most heavily for development support in the parts of the project that were changing the fastest.",
        "The work focused on fixing bugs, improving reliability, and connecting frontend and backend behavior.",
        "The biggest usage appears in multi-file feature areas rather than isolated one-line edits.",
        "Examples include commits such as 0ab49a8 feat(api): improve library and workspace workflows and c668c88 feat(ui): improve study hub experience.",
    ]),
    ("Admin Interface", [
        "AI helped refine the admin UI and supporting logic.",
        "This included moderation-related screens, admin controls, and activity log behavior.",
        "The changes suggest AI was used to smooth out page behavior, improve data display, and make admin workflows more dependable.",
        "The admin work is visible in 62ff978 feat: refine admin UI and notifications, which touched ActivityLogPage, AdminLayout, AdminProfilePage, NotificationsPage, SettingPage, LibraryPage, AIContentModerationPage, and WorkSpacePage.",
        "Earlier admin-focused work also appears in 3e35cec feat(admin-logs): restyle audit log filters, 1f46e74 feat(admin-users): improve user management controls, 2f66907 feat(admin-moderation): refine filters and layout, and ba38637 feat(admin-dashboard): visualize admin metrics.",
    ]),
    ("AI Moderation Workflow", [
        "The moderation flow appears to have required repeated adjustments.",
        "AI was likely used to help handle case display, moderation actions, and review-page behavior.",
        "Backend service updates also suggest AI helped align moderation data with what the frontend expected.",
        "The moderation page was part of the 62ff978 admin refinement pass, and later fixes continued in 8319320 fix all and 98922f1 Fix handle file and admin page.",
    ]),
    ("Activity Log Area", [
        "AI was used to improve how logs were shown and handled.",
        "This likely covered filtering, formatting, and consistency of displayed records.",
        "The logging backend was also adjusted so the page could reflect accurate activity data.",
        "The history shows a specific admin-log pass in 3e35cec feat(admin-logs): restyle audit log filters and again in 62ff978, where ActivityLogPage.css and ActivityLogPage.jsx were both updated.",
    ]),
    ("Document Upload and File Handling", [
        "This is one of the strongest AI usage areas.",
        "AI was used to help fix document upload reliability and file-processing behavior.",
        "Tag suggestion logic and duplicate-upload prevention were also improved.",
        "The work here looks like a mix of debugging, validation, and workflow repair.",
        "The clearest examples are 670a546 Improve document uploads, tagging, and reliability and d1debb6 Prevent duplicate document uploads.",
        "Those commits changed documentController.js, aiService.js, and LibraryPage.jsx, showing that the upload flow required coordination across backend logic and the library UI.",
        "98922f1 Fix handle file and admin page continued that line of work by updating documentController.js, aiService.js, LibraryPage.jsx, and WorkSpacePage.jsx together.",
    ]),
    ("Library Page", [
        "AI helped improve how library content was displayed and interacted with.",
        "This area seems tied to document browsing and upload results.",
        "It likely needed coordination with the document backend so the page showed the right items and metadata.",
        "The library work shows up repeatedly in 0ab49a8, c668c88, 670a546, 98922f1, and 6e50869 Fix Admin and LibraryPage.",
    ]),
    ("Workspace Page", [
        "The workspace received large-scale changes, which points to substantial AI involvement.",
        "AI was likely used to repair broken interactions, simplify complex logic, and stabilize the page layout and flow.",
        "The size of the changes suggests this was one of the most difficult parts of the project.",
        "The biggest workspace refactor appears in 98922f1, where WorkSpacePage.jsx alone changed by thousands of lines, and in c668c88, where WorkSpacePage.css and WorkSpacePage.jsx were both reworked.",
        "Other workspace-related lines include 8b82f83 Improve workspace and library interactions, bfa1b21 Improve workspace solution experience, and 50b8b1a Add contributor solution permissions.",
    ]),
    ("Backend Controllers and Services", [
        "AI support extended beyond the frontend into backend logic.",
        "adminController.js, documentController.js, activityLogService.js, and aiService.js all saw meaningful updates.",
        "That indicates AI was being used not just for appearance fixes, but for end-to-end feature behavior.",
        "The API/workflow line is especially visible in 0ab49a8 feat(api): improve library and workspace workflows, which changed aiController.js, documentController.js, workspaceController.js, documentRoutes.js, workspaceRoutes.js, and aiService.js.",
        "Later backend work continued in 8319320 fix all, which touched adminController.js, documentController.js, adminRoutes.js, activityLogService.js, aiService.js, and FE/src/utils/adminApi.js.",
    ]),
    ("Solution Comment Support", [
        "The backend gained or improved support for solution comments.",
        "This suggests AI was used to add or repair collaboration-related functionality.",
        "It likely helped connect comment features to existing document or workspace workflows.",
        "A direct example is 1f54b04 Add solution comment backend support.",
        "That was followed by bfa1b21 Improve workspace solution experience and 50b8b1a Add contributor solution permissions, which suggests the comment workflow was expanded in stages.",
    ]),
    ("Project-Wide Cleanup", [
        "Several commits labeled as broad fixes suggest AI was used for general stabilization.",
        "This included cleaning up broken behavior across multiple pages and features.",
        "The language change to English also indicates a final polish step for consistency.",
        "The history includes several broad cleanup passes such as 2090bef fix all, e79a140 fix all, and 8319320 fix all.",
        "The final commit in the visible history, 04ea0eb Change language to english, shows the project was also normalized for English presentation at the end.",
    ]),
    ("Overall Usage Pattern", [
        "AI was used mainly as a coding assistant for debugging, refactoring, feature repair, backend/frontend alignment, and workflow stabilization.",
        "The strongest benefit seems to have been in complex areas where one change needed to be coordinated across many files.",
        "The pattern across commit history suggests the AI was most useful when the project needed coordinated edits to UI, APIs, services, and route logic at the same time.",
    ]),
    ("What Is Not Measured", [
        "This report does not include token counts, the number of prompts, exact session duration, or billing data.",
        "It reflects project history and change patterns rather than raw AI telemetry.",
    ]),
]

for heading, bullets in sections:
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    p.add_run(heading)
    if bullets:
        for item in bullets:
            b = doc.add_paragraph(style="Normal")
            b.style = styles["Normal"]
            b.paragraph_format.left_indent = Inches(0.25)
            b.paragraph_format.space_after = Pt(4)
            b.add_run("• ").bold = True
            b.add_run(item)

table_heading = doc.add_paragraph(style="Heading 2")
table_heading.paragraph_format.space_before = Pt(12)
table_heading.paragraph_format.space_after = Pt(6)
table_heading.add_run("Usage Summary")

table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"
table.autofit = False
table.columns[0].width = Inches(1.8)
table.columns[1].width = Inches(3.0)
table.columns[2].width = Inches(2.5)
hdr = table.rows[0].cells
headers = ["Area", "AI Involvement", "Impact"]
for cell, text in zip(hdr, headers):
    cell.text = text
    set_cell_shading(cell, "D9EAF7")
    set_cell_margins(cell)
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.bold = True
            r.font.name = "Arial"
            r.font.size = Pt(10.5)

rows = [
    ("Admin", "High", "Moderation, activity log, and permission fixes"),
    ("Documents", "High", "Upload reliability, tagging, duplicate prevention"),
    ("Workspace", "High", "Large refactors and interaction repair"),
    ("Backend", "Medium-High", "Controller and AI service alignment"),
    ("Project-wide cleanup", "Medium", "Broad stabilization and language normalization"),
]

for area, involvement, impact in rows:
    row = table.add_row().cells
    for idx, text in enumerate((area, involvement, impact)):
        row[idx].text = text
        set_cell_margins(row[idx])
        for p in row[idx].paragraphs:
            for r in p.runs:
                r.font.name = "Arial"
                r.font.size = Pt(10.5)

doc.add_paragraph(
    "Prepared for project reporting purposes. The analysis is based on repository history and visible change patterns.",
    style="Normal",
)

doc.save(OUT_FILE)
